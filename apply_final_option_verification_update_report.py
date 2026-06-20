from pathlib import Path
import shutil

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path("/Users/castle/C盘/常用/以往实习梳理/易方达产品")
BACKUP_DIR = ROOT / "备份文件"
REPORT_PATH = ROOT / "海外ETF期权及衍生策略产品研究_正文.docx"
WORKBOOK_PATH = ROOT / "海外ETF期权及衍生策略产品研究_核验底稿.xlsx"
FINAL_VERIFY_PATH = Path("/Users/castle/Downloads/期权使用核验结果_二级复核后最终简表.xlsx")
REPORT_BACKUP = BACKUP_DIR / "海外ETF期权及衍生策略产品研究_正文_期权最终核验同步前备份_20260617.docx"
WORKBOOK_BACKUP = BACKUP_DIR / "海外ETF期权及衍生策略产品研究_核验底稿_期权最终核验同步前备份_20260617.xlsx"
MD_PATH = ROOT / "衍生策略产品分类与期权使用识别口径说明.md"

CN_FONT = "KaiTi"
EN_FONT = "Times New Roman"
INK = "000000"
GREY = "6B7280"
HEADER_FILL = "F2F4F7"
BORDER = "D0D7DE"
YELLOW = "FFF2CC"

REGION_ORDER = ["美国", "欧洲", "亚太", "加拿大", "其他"]
TYPE_ORDER = ["收益增强型", "风险缓冲型", "杠杆与反向型", "另类衍生策略型"]


def clean(value):
    if pd.isna(value):
        return ""
    return str(value).strip()


def money_yi(value, digits=2):
    return f"{value / 1e8:,.{digits}f}"


def pct(value):
    return f"{value * 100:.2f}%"


def font_run(run, size=None, bold=None, color=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = EN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), EN_FONT)
    r_fonts.set(qn("w:hAnsi"), EN_FONT)
    r_fonts.set(qn("w:eastAsia"), CN_FONT)


def set_para_format(paragraph, size=11, before=0, after=6, line=1.2, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        font_run(run, size=size)


def add_before(doc, ref, text, style_name="Normal", size=11, bold=False, color=INK, before=0, after=6, align=None):
    paragraph = doc.add_paragraph(style=style_name)
    if text:
        run = paragraph.add_run(text)
        font_run(run, size=size, bold=bold, color=color)
    set_para_format(paragraph, size=size, before=before, after=after, align=align)
    ref.addprevious(paragraph._element)
    return paragraph


def add_heading_before(doc, ref, text, level):
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    size = {1: 16, 2: 13, 3: 12}[level]
    before = {1: 16, 2: 12, 3: 8}[level]
    after = {1: 8, 2: 6, 3: 4}[level]
    return add_before(doc, ref, text, style_name, size=size, bold=True, color=INK, before=before, after=after)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_autofit_window(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is not None:
        tbl_pr.remove(tbl_layout)
    set_table_borders(table)


def fill_cell(cell, text, style_name, size=8.2, bold=False, shade=None):
    if shade:
        set_cell_shading(cell, shade)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    try:
        p.style = style_name
    except KeyError:
        pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("" if text is None else str(text))
    font_run(run, size=size, bold=bold, color=INK)


def add_doc_table(doc, ref, headers, rows, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_autofit_window(table)
    for idx, header in enumerate(headers):
        fill_cell(table.rows[0].cells[idx], header, "表头", size=8.8, bold=True, shade=HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            text = "" if value is None else str(value)
            shade = YELLOW if "待补充" in text else None
            fill_cell(cells[idx], text, "表格正文", size=font_size, shade=shade)
    ref.addprevious(table._element)
    return table


def remove_section_after_heading(doc, start_text, end_text):
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = None
    end_idx = None
    for idx, child in enumerate(children):
        if isinstance(child, CT_P):
            text = "".join(child.itertext()).strip()
            if start_idx is None and text.startswith(start_text):
                start_idx = idx
            elif start_idx is not None and text.startswith(end_text):
                end_idx = idx
                break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Cannot find section range: {start_text} to {end_text}")
    ref = children[end_idx]
    for child in children[start_idx + 1 : end_idx]:
        body.remove(child)
    return ref


def remove_section_with_heading(doc, start_text, end_text):
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = None
    end_idx = None
    for idx, child in enumerate(children):
        if isinstance(child, CT_P):
            text = "".join(child.itertext()).strip()
            if start_idx is None and text.startswith(start_text):
                start_idx = idx
            elif start_idx is not None and text.startswith(end_text):
                end_idx = idx
                break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Cannot find section range: {start_text} to {end_text}")
    ref = children[end_idx]
    for child in children[start_idx:end_idx]:
        body.remove(child)
    return ref


def style_ws(ws):
    thin = Side(style="thin", color=BORDER)
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.font = Font(name=CN_FONT, size=10)
            cell.border = border
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor=HEADER_FILL)
        cell.font = Font(name=CN_FONT, size=10, bold=True)


def set_widths(ws, widths):
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width


def strategy_type(row, product_type):
    text = " ".join([clean(row.get("判断说明")), clean(row.get("产品名称"))]).lower()
    if product_type == "收益增强型":
        return "收益增强"
    if product_type == "风险缓冲型":
        return "风险缓冲"
    if "collar" in text or "collared" in text:
        return "期权collar"
    if "tail risk" in text or "protective put" in text or "put option" in text:
        return "尾部风险或保护性期权"
    if "box spread" in text:
        return "期权价差或box spread"
    if "overlay" in text or "hedg" in text:
        return "期权对冲或叠加"
    return "其他期权策略"


def load_and_apply_final():
    final = pd.read_excel(FINAL_VERIFY_PATH)
    final["SecId"] = final["SecId"].map(clean)
    final["最终判断"] = final["最终判断"].map(clean)
    if set(final["最终判断"].unique()) - {"是", "否"}:
        raise RuntimeError("最终判断列存在非是/否值")

    wb = load_workbook(WORKBOOK_PATH)
    detail_ws = wb["衍生策略分类明细"]
    headers = {cell.value: idx + 1 for idx, cell in enumerate(detail_ws[1])}
    final_map = {row["SecId"]: row for _, row in final.iterrows()}

    for row_idx in range(2, detail_ws.max_row + 1):
        secid = clean(detail_ws.cell(row_idx, headers["SecId"]).value)
        if secid not in final_map:
            continue
        result = final_map[secid]
        conclusion = result["最终判断"]
        product_type = clean(detail_ws.cell(row_idx, headers["产品类型"]).value)
        detail_ws.cell(row_idx, headers["是否使用期权"]).value = conclusion
        if conclusion == "是":
            detail_ws.cell(row_idx, headers["期权策略类型"]).value = strategy_type(result, product_type)
            detail_ws.cell(row_idx, headers["期权识别依据"]).value = f"二级复核确认使用期权：{clean(result.get('判断说明'))}"
        else:
            detail_ws.cell(row_idx, headers["期权策略类型"]).value = "未使用期权的衍生策略"
            detail_ws.cell(row_idx, headers["期权识别依据"]).value = f"二级复核确认未使用期权：{clean(result.get('判断说明'))}"
    style_ws(detail_ws)

    # Re-read updated detail from worksheet values.
    detail_df = pd.DataFrame(detail_ws.values)
    detail_df.columns = detail_df.iloc[0]
    detail_df = detail_df.iloc[1:].copy()
    detail_df["Fund Size USD"] = pd.to_numeric(detail_df["Fund Size USD"], errors="coerce").fillna(0)

    # Replace final verification sheet.
    sheet_name = "期权使用二级复核结果"
    if sheet_name in wb.sheetnames:
        del wb[sheet_name]
    ws_final = wb.create_sheet(sheet_name)
    out_cols = ["产品名称", "Ticker", "SecId", "AUM（USD）", "最终判断", "信源", "置信度", "判断说明"]
    ws_final.append(out_cols)
    for _, row in final.iterrows():
        ws_final.append([clean(row.get(col)) for col in out_cols])
    style_ws(ws_final)
    set_widths(ws_final, [36, 12, 14, 14, 10, 48, 10, 60])
    ws_final.freeze_panes = "A2"
    ws_final.auto_filter.ref = ws_final.dimensions

    # Replace manual verification list with final simple table.
    ws_review = wb["期权使用人工核验清单"]
    ws_review.delete_rows(1, ws_review.max_row)
    ws_review.append(["产品名称", "Ticker", "SecId", "AUM（USD）", "人工核验是否使用期权", "信源", "置信度", "判断说明"])
    for _, row in final.iterrows():
        ws_review.append(
            [
                clean(row.get("产品名称")),
                clean(row.get("Ticker")),
                clean(row.get("SecId")),
                row.get("AUM（USD）"),
                clean(row.get("最终判断")),
                clean(row.get("信源")),
                clean(row.get("置信度")),
                clean(row.get("判断说明")),
            ]
        )
    style_ws(ws_review)
    set_widths(ws_review, [36, 12, 14, 14, 16, 48, 10, 60])
    ws_review.freeze_panes = "A2"
    ws_review.auto_filter.ref = ws_review.dimensions

    # Update option usage stats.
    stat_ws = wb["期权使用识别统计"]
    stat_ws.delete_rows(2, stat_ws.max_row)
    total_count = len(detail_df)
    total_aum = detail_df["Fund Size USD"].sum()
    for flag in ["是", "否"]:
        sub = detail_df[detail_df["是否使用期权"].eq(flag)]
        aum = sub["Fund Size USD"].sum()
        note = "规则识别并已完成二级人工复核；表示产品本身、底层基金或策略直接或间接使用期权" if flag == "是" else "规则识别并已完成二级人工复核；表示已判断其衍生策略不依赖期权"
        stat_ws.append([flag, int(sub["SecId"].count()), sub["SecId"].count() / total_count, aum, aum / 1e8, aum / total_aum, note])
    style_ws(stat_ws)
    for row_idx in range(2, stat_ws.max_row + 1):
        stat_ws.cell(row_idx, 3).number_format = "0.00%"
        stat_ws.cell(row_idx, 6).number_format = "0.00%"
        stat_ws.cell(row_idx, 4).number_format = "#,##0.00"
        stat_ws.cell(row_idx, 5).number_format = "#,##0.00"

    # Update regional stats sheet.
    reg_ws = wb["1.2区域统计"]
    reg_headers = {cell.value: idx + 1 for idx, cell in enumerate(reg_ws[1])}
    option_df = detail_df[detail_df["是否使用期权"].eq("是")]
    for row_idx in range(2, reg_ws.max_row + 1):
        region = reg_ws.cell(row_idx, reg_headers["地区"]).value
        global_aum = float(reg_ws.cell(row_idx, reg_headers["全球ETF AUM（USD）"]).value or 0)
        if region == "合计":
            sub = option_df
        else:
            sub = option_df[option_df["Region"].eq(region)]
        count = int(sub["SecId"].count())
        aum = float(sub["Fund Size USD"].sum())
        reg_ws.cell(row_idx, reg_headers["使用期权产品数量"]).value = count
        reg_ws.cell(row_idx, reg_headers["使用期权产品AUM（USD）"]).value = aum
        reg_ws.cell(row_idx, reg_headers["使用期权产品AUM（亿美元）"]).value = aum / 1e8
        reg_ws.cell(row_idx, reg_headers["使用期权产品规模渗透率"]).value = aum / global_aum if global_aum else 0
    style_ws(reg_ws)

    # Update product structure sheet.
    struct_ws = wb["1.3产品结构统计"]
    struct_ws.delete_rows(1, struct_ws.max_row)
    struct_headers = [
        "产品类型",
        "产品数量",
        "数量占比",
        "AUM（USD）",
        "AUM（亿美元）",
        "AUM占比",
        "使用期权产品数量",
        "使用期权AUM（USD）",
        "使用期权AUM（亿美元）",
        "使用期权AUM占该类比例",
        "分类说明",
        "代表产品",
    ]
    struct_ws.append(struct_headers)
    type_desc = {
        "收益增强型": ("以权利金、期权收入和现金流分派为主", "JEPI、JEPQ、QYLD"),
        "风险缓冲型": ("通过FLEX Options或期权组合设定buffer和cap", "BUFR、BALT、PJAN"),
        "杠杆与反向型": ("主要通过期货、互换和每日再平衡实现，少数产品使用期权", "TQQQ、SOXL、SSO"),
        "另类衍生策略型": ("覆盖尾部风险、波动率、管理期货、多资产叠加等，工具需逐项识别", "IALT、DBMF、HELO、ACIO"),
    }
    for typ in TYPE_ORDER:
        sub = detail_df[detail_df["产品类型"].eq(typ)]
        opt = sub[sub["是否使用期权"].eq("是")]
        aum = sub["Fund Size USD"].sum()
        opt_aum = opt["Fund Size USD"].sum()
        struct_ws.append(
            [
                typ,
                int(sub["SecId"].count()),
                sub["SecId"].count() / total_count,
                aum,
                aum / 1e8,
                aum / total_aum,
                int(opt["SecId"].count()),
                opt_aum,
                opt_aum / 1e8,
                opt_aum / aum if aum else 0,
                type_desc[typ][0],
                type_desc[typ][1],
            ]
        )
    style_ws(struct_ws)

    # Source and claim records.
    source_ws = wb["来源清单"]
    titles = {clean(source_ws.cell(row, 3).value): row for row in range(2, source_ws.max_row + 1)}
    source_title = FINAL_VERIFY_PATH.name
    if source_title not in titles:
        source_id = f"S{source_ws.max_row:02d}"
        source_ws.append([source_id, "Excel", source_title, str(FINAL_VERIFY_PATH), "2026-06-17", "是否使用期权识别、1.2、1.3、第二章管理人格局", "用户二级复核最终简表"])
    else:
        source_id = clean(source_ws.cell(titles[source_title], 1).value)
    style_ws(source_ws)

    claim_ws = wb["事实核验表"]
    claims = {clean(claim_ws.cell(row, 3).value) for row in range(2, claim_ws.max_row + 1)}
    claim_text = "二级复核最终简表显示，368只待核验及灰区产品中74只确认为使用期权，294只确认为未使用期权"
    if claim_text not in claims:
        existing_ids = []
        for row in range(2, claim_ws.max_row + 1):
            val = clean(claim_ws.cell(row, 1).value)
            if val.startswith("C") and val[1:].isdigit():
                existing_ids.append(int(val[1:]))
        claim_id = f"C{max(existing_ids or [0]) + 1:03d}"
        claim_ws.append([claim_id, "期权使用识别统计", claim_text, source_id, "期权使用核验结果_二级复核后最终简表.xlsx：最终判断列", "已按SecId同步", "同步后底稿已无待人工核验产品"])
    style_ws(claim_ws)

    # Classification note.
    note_ws = wb["分类口径说明"]
    for row_idx in range(2, note_ws.max_row + 1):
        if note_ws.cell(row_idx, 1).value in {"期权使用识别", "是否使用期权"}:
            note_ws.cell(row_idx, 3).value = "基于Morningstar字段规则识别，并叠加灰区关键词复核及二级复核最终简表；当前样本已无待人工核验产品"
    style_ws(note_ws)

    wb.save(WORKBOOK_PATH)

    # Return updated dataframes and stats.
    detail_df["Fund Size USD"] = pd.to_numeric(detail_df["Fund Size USD"], errors="coerce").fillna(0)
    return detail_df


def build_stats(detail):
    total_aum = detail["Fund Size USD"].sum()
    opt = detail[detail["是否使用期权"].eq("是")]
    non = detail[detail["是否使用期权"].eq("否")]
    reg = pd.read_excel(WORKBOOK_PATH, sheet_name="1.2区域统计")
    region_rows = []
    for _, row in reg.iterrows():
        region_rows.append(
            [
                row["地区"],
                f"{row['全球ETF AUM（亿美元）']:,.2f}",
                f"{int(row['衍生策略产品数量']):,}",
                f"{row['衍生策略AUM（亿美元）']:,.2f}",
                pct(row["规模渗透率"]),
                f"{int(row['使用期权产品数量']):,}",
                f"{row['使用期权产品AUM（亿美元）']:,.2f}",
                pct(row["使用期权产品规模渗透率"]),
            ]
        )

    type_rows = []
    cross_count = pd.crosstab(detail["产品类型"], detail["是否使用期权"])
    cross_aum = pd.pivot_table(detail, index="产品类型", columns="是否使用期权", values="Fund Size USD", aggfunc="sum", fill_value=0)
    type_labels = {
        "收益增强型": "收益增强",
        "风险缓冲型": "风险缓冲",
        "杠杆与反向型": "杠杆与反向",
        "另类衍生策略型": "另类衍生策略",
    }
    for typ in TYPE_ORDER:
        sub = detail[detail["产品类型"].eq(typ)]
        aum = sub["Fund Size USD"].sum()
        opt_count = int(cross_count.loc[typ, "是"]) if "是" in cross_count.columns else 0
        opt_aum = float(cross_aum.loc[typ, "是"]) if "是" in cross_aum.columns else 0
        if typ == "收益增强型":
            note = "JEPI、JEPQ、QYLD；权利金和分派属性突出"
        elif typ == "风险缓冲型":
            note = "BUFR、BALT、PJAN；通过期权组合设定保护区间和收益上限"
        elif typ == "杠杆与反向型":
            note = "TQQQ、SOXL、SSO；绝大多数通过期货、互换和每日再平衡实现"
        else:
            note = "IALT、DBMF、HELO、ACIO；工具差异较大，已完成是否使用期权复核"
        type_rows.append(
            [
                type_labels[typ],
                f"{len(sub):,}",
                pct(len(sub) / len(detail)),
                f"{aum / 1e8:,.2f}",
                pct(aum / total_aum),
                f"{opt_count:,}",
                f"{opt_aum / 1e8:,.2f}",
                pct(opt_aum / aum if aum else 0),
                note,
            ]
        )
    type_rows.append(["合计", f"{len(detail):,}", "100.00%", f"{total_aum / 1e8:,.2f}", "100.00%", f"{len(opt):,}", f"{opt['Fund Size USD'].sum() / 1e8:,.2f}", pct(opt["Fund Size USD"].sum() / total_aum), "-"])

    manager = opt.groupby("Firm Name").agg(产品数量=("SecId", "count"), AUM=("Fund Size USD", "sum")).reset_index()
    manager["AUM占比"] = manager["AUM"] / opt["Fund Size USD"].sum()
    manager = manager.sort_values("AUM", ascending=False)

    return {
        "total_aum": total_aum,
        "opt": opt,
        "non": non,
        "region_rows": region_rows,
        "type_rows": type_rows,
        "manager": manager,
        "cross_count": cross_count,
        "cross_aum": cross_aum,
    }


def top_products_text(detail, typ):
    sub = detail[detail["产品类型"].eq(typ)].sort_values("Fund Size USD", ascending=False).head(5)
    total = detail.loc[detail["产品类型"].eq(typ), "Fund Size USD"].sum()
    names = "、".join([f"{row['Name']}（{row['Fund Size USD'] / 1e8:.2f}亿美元）" for _, row in sub.iterrows()])
    return names, sub["Fund Size USD"].sum() / 1e8, sub["Fund Size USD"].sum() / total


def update_report(detail, stats):
    doc = Document(REPORT_PATH)

    # 1.2
    ref = remove_section_after_heading(doc, "1.2市场规模与区域分布", "1.3全球衍生策略产品结构总览")
    add_before(doc, ref, "根据Morningstar数据，截至2026年5月31日，全球ETF AUM合计约21.35万亿美元。从基金注册地看，美国AUM约15.61万亿美元，占全球73.14%；欧洲约3.71万亿美元，占17.38%；亚太约1.47万亿美元，占6.87%；加拿大约0.51万亿美元，占2.39%。全球ETF资产仍高度集中在美国市场，这也是后续衍生策略产品最早形成规模化供给的基础。")
    add_before(doc, ref, "表：全球ETF与衍生策略产品区域分布", "表格标题", size=10.5, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_doc_table(
        doc,
        ref,
        ["地区", "全球ETF AUM（亿美元）", "衍生策略产品数量", "衍生策略AUM（亿美元）", "衍生策略规模渗透率", "使用期权产品数量", "使用期权AUM（亿美元）", "使用期权规模渗透率"],
        stats["region_rows"],
        font_size=7.5,
    )
    add_before(doc, ref, "资料来源：Morningstar、易方达产品研究", "资料来源", size=9, color=GREY, after=2)
    add_before(doc, ref, "注：全球ETF AUM按用户设定的Morningstar数据口径，截至2026年5月31日；衍生策略产品数量按当前导出存续样本统计，AUM使用导出文件中的Fund Size USD字段；原表中的Fund Size Date为单产品规模日期字段，存在晚于2026年5月31日的记录，不作为本报告整体统计截止口径；如需严格锁定2026年5月31日AUM，需另行导出历史时点Fund Size数据；规模渗透率为对应产品AUM除以同地区全球ETF AUM；数量为样本内产品数量，不计算数量渗透率", "注释", size=9, color=GREY, after=6)

    opt = stats["opt"]
    non = stats["non"]
    total_aum = stats["total_aum"]
    add_before(doc, ref, f"进一步看衍生策略产品，根据Morningstar数据，截至2026年5月31日，全球样本共3,161只，AUM约{total_aum / 1e8:.2f}亿美元，约占全球ETF AUM的2.96%。其中，使用期权的衍生策略产品共{len(opt):,}只，AUM约{opt['Fund Size USD'].sum() / 1e8:.2f}亿美元，占衍生策略样本AUM的{opt['Fund Size USD'].sum() / total_aum * 100:.2f}%；未使用期权的衍生策略产品共{len(non):,}只，AUM约{non['Fund Size USD'].sum() / 1e8:.2f}亿美元。")
    add_before(doc, ref, f"美国仍是使用期权产品的核心市场，样本数量为{int(opt[opt['Region'].eq('美国')]['SecId'].count()):,}只，AUM约{opt[opt['Region'].eq('美国')]['Fund Size USD'].sum() / 1e8:.2f}亿美元，占全球使用期权产品AUM的{opt[opt['Region'].eq('美国')]['Fund Size USD'].sum() / opt['Fund Size USD'].sum() * 100:.2f}%。加拿大使用期权产品AUM约{opt[opt['Region'].eq('加拿大')]['Fund Size USD'].sum() / 1e8:.2f}亿美元，主要体现为备兑认购和收益型产品；欧洲和亚太市场虽然也有相关产品，但规模占比明显低于美国。")
    add_before(doc, ref, "区域差异上，亚太衍生策略产品规模渗透率较高，主要与韩国、日本、中国台湾和中国香港等市场中杠杆与反向产品较多有关；但从使用期权产品看，美国和加拿大更具代表性。欧洲市场虽然ETF总规模较大，但在UCITS、PRIIPs KID和ESMA风险计量框架下，复杂衍生策略产品供给更偏审慎。")

    # 1.3
    ref = remove_section_after_heading(doc, "1.3全球衍生策略产品结构总览", "二、头部管理人")
    add_before(doc, ref, "按照产品功能和收益结构，全球衍生策略产品可以归为收益增强型、风险缓冲型、杠杆与反向型、另类衍生策略型。从数量看，杠杆与反向型产品最多，共1,843只，占样本58.30%；风险缓冲型520只，占16.45%；另类衍生策略型468只，占14.81%；收益增强型330只，占10.44%。")
    add_before(doc, ref, "从是否使用期权看，收益增强型和风险缓冲型均属于使用期权的产品；杠杆与反向型绝大多数并不使用期权，主要通过期货、互换和每日再平衡实现；另类衍生策略型内部差异最大，既包括尾部风险、collar、options overlay等期权策略，也包括管理期货、市场中性、货币对冲和普通指数跟踪产品。本次二级复核后，样本已无待人工核验产品。")
    names, aum, share = top_products_text(detail, "收益增强型")
    add_before(doc, ref, f"收益增强型前五大产品为{names}，合计规模约{aum:.2f}亿美元，占该类AUM的{share * 100:.2f}%。")
    names, aum, share = top_products_text(detail, "风险缓冲型")
    add_before(doc, ref, f"风险缓冲型前五大产品为{names}，合计规模约{aum:.2f}亿美元，占该类AUM的{share * 100:.2f}%。")
    names, aum, share = top_products_text(detail, "杠杆与反向型")
    add_before(doc, ref, f"杠杆与反向型前五大产品为{names}，合计规模约{aum:.2f}亿美元，占该类AUM的{share * 100:.2f}%。这类产品规模较大，但通常不属于期权策略产品。")
    names, aum, share = top_products_text(detail, "另类衍生策略型")
    add_before(doc, ref, f"另类衍生策略型前五大产品为{names}，合计规模约{aum:.2f}亿美元，占该类AUM的{share * 100:.2f}%。该类产品不能仅凭产品名称判断是否使用期权，需要结合底层策略、持仓和招募说明书逐项识别。")
    add_before(doc, ref, "表：全球衍生策略产品结构分布", "表格标题", size=10.5, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_doc_table(
        doc,
        ref,
        ["产品类型", "数量", "数量占比", "AUM（亿美元）", "AUM占比", "使用期权产品数量", "使用期权AUM（亿美元）", "使用期权AUM占该类比例", "代表产品与说明"],
        stats["type_rows"],
        font_size=7.4,
    )
    add_before(doc, ref, "资料来源：Morningstar、管理人产品资料、易方达产品研究", "资料来源", size=9, color=GREY, after=6)

    # Chapter 2
    ref = remove_section_with_heading(doc, "二、头部管理人", "三、主要产品类型拆解")
    add_heading_before(doc, ref, "二、使用期权产品管理人布局与竞争优势", 1)
    manager = stats["manager"]
    top5 = manager.head(5)
    add_before(doc, ref, "本章聚焦使用期权的衍生策略产品管理人，而不是按全部衍生策略产品口径排序。按AUM看，前五大管理人为JPMorgan、First Trust、Innovator ETFs、NEOS和Global X。它们的竞争优势并不相同：JPMorgan依靠主动权益组合、ELN执行和渠道能力形成大单品；First Trust和Innovator依靠Buffer产品货架和目标结果教育形成规模；NEOS和Global X则分别在高收入叙事和规则化buy-write复制上切入市场。")
    add_before(doc, ref, "表：使用期权的衍生策略产品头部管理人", "表格标题", size=10.5, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = []
    features = {
        "JPMorgan": "主动权益组合叠加ELN和期权覆盖",
        "First Trust": "Buffer货架与Laddered组合",
        "Innovator ETFs": "Defined Outcome与Buffer系列化扩张",
        "Neos Funds": "高收入option income与月度分派",
        "Global X Funds": "规则化buy-write和风险管理收入产品",
    }
    for idx, row in enumerate(top5.itertuples(index=False), start=1):
        rows.append([idx, row._0, f"{int(row.产品数量):,}", f"{row.AUM / 1e8:.2f}", pct(row.AUM占比), features.get(row._0, "")])
    add_doc_table(doc, ref, ["排名", "管理人", "产品数量", "AUM（亿美元）", "占使用期权产品AUM", "布局特征"], rows, font_size=7.6)
    add_before(doc, ref, "资料来源：Morningstar、管理人产品资料、易方达产品研究", "资料来源", size=9, color=GREY, after=6)

    add_heading_before(doc, ref, "2.1 J.P. Morgan：主动权益收益大单品路径", 2)
    add_before(doc, ref, "J.P. Morgan在使用期权的衍生策略产品中AUM约896.49亿美元，样本内产品数量为7只，排名第一。其核心路径是先通过主动管理型权益收益产品建立规模。JEPI于2020年5月20日发行，JEPQ于2022年5月3日发行，两只产品的年化费用均为0.35%。两只基金均强调在保留资本增值机会的同时获取当期收入，并通过权益组合和期权相关结构实现月度分派。")
    add_before(doc, ref, "JEPI的权益组合面向美国大盘股票，并通过ELN嵌入对S&P 500 Total Return Index或复制该指数的ETF的卖出认购期权敞口。JEPQ的权益组合显著面向Nasdaq 100相关股票，并通过ELN嵌入对Nasdaq 100的卖出认购期权敞口。两只产品的优势不在于机械复制指数并卖出期权，而在于主动选股、较低费率、月度分派叙事和财富管理渠道共同形成的大单品效应。")
    add_before(doc, ref, "表：J.P. Morgan收益增强产品对比", "表格标题", size=10.5, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_doc_table(doc, ref, ["产品", "发行时间", "费率", "底层资产", "期权工具", "规模"], [["JEPI", "2020-05-20", "0.35%", "美国大盘股票", "ELN中嵌入卖出认购期权", "445.48亿美元"], ["JEPQ", "2022-05-03", "0.35%", "Nasdaq 100相关股票", "ELN中嵌入卖出认购期权", "401.53亿美元"]])
    add_before(doc, ref, "资料来源：J.P. Morgan、Morningstar、易方达产品研究", "资料来源", size=9, color=GREY, after=6)

    add_heading_before(doc, ref, "2.2 First Trust：Buffer货架与组合化产品", 2)
    add_before(doc, ref, "First Trust在使用期权的衍生策略产品中AUM约594.26亿美元，样本内产品数量为138只。其核心路径是把单期Buffer结构做成可持续滚动的产品货架，并进一步通过Laddered产品降低单一建仓时点的影响。First Trust Canada BUFR的资料显示，该基金通过等权投资四只底层Target Outcome Buffer ETF，获得对S&P 500相关美国大盘股票的缓冲型敞口；每三个月有一只底层ETF进入新的为期一年的目标结果期间，相应刷新buffer和cap。")
    add_before(doc, ref, "这一设计把原本需要投资者自行选择月份和建仓时点的Buffer产品，转化为更容易长期持有的组合工具。竞争优势来自产品矩阵、结构标准化和渠道教育，而不是单只基金的短期收益表现。")
    add_before(doc, ref, "表：First Trust Buffer产品布局", "表格标题", size=10.5, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_doc_table(doc, ref, ["产品", "市场", "发行时间", "管理费", "底层结构", "策略定位"], [["BUFR", "加拿大", "2023-05-25", "0.15%", "四只FT Vest Target Outcome Buffer ETF", "Laddered Buffer配置工具"], ["FT Vest Buffer系列", "美国与加拿大", "待补充最新数据", "待补充最新数据", "不同月份和保护区间的Buffer产品", "Buffer产品货架"]])
    add_before(doc, ref, "资料来源：First Trust、Morningstar、易方达产品研究", "资料来源", size=9, color=GREY, after=6)

    add_heading_before(doc, ref, "2.3 Innovator：目标结果品类教育与系列化扩张", 2)
    add_before(doc, ref, "Innovator ETFs在使用期权的衍生策略产品中AUM约314.87亿美元，样本内产品数量为165只。其竞争优势在于目标结果产品的品类教育。其产品表显示，Innovator已经形成美国股票Buffer、Power Buffer、Ultra Buffer、Growth-100 Power Buffer、小盘股、国际发达市场、新兴市场、季度Buffer、100% Buffer、收入型、加速型、管理型Buffer和Managed Outcome等多个系列。产品扩张逻辑非常清晰：先定义保护区间、收益上限和outcome period，再按底层资产、月份和保护强度横向复制。")
    add_before(doc, ref, "与First Trust相比，Innovator更强调产品风险收益区间的可解释性。对国内产品开发的启示是：若未来做缓冲型产品，投资者教育必须前置，尤其要明确buffer不是保本，中途买入的剩余buffer和cap会变化，超过缓冲区间的亏损仍由投资者承担。")
    add_before(doc, ref, "表：Innovator目标结果产品系列", "表格标题", size=10.5, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_doc_table(doc, ref, ["系列", "保护区间", "Outcome Period", "底层资产", "说明"], [["U.S. Equity Buffer ETFs", "9% Buffer", "1年", "S&P 500 ETF", "基础缓冲系列"], ["U.S. Equity Power Buffer ETFs", "15% Buffer", "1年", "S&P 500 ETF", "更高缓冲系列"], ["U.S. Equity Ultra Buffer ETFs", "30% Buffer，区间为下跌5%至35%", "1年", "S&P 500 ETF", "区间式缓冲"], ["Growth-100 Power Buffer ETFs", "15% Buffer", "1年", "Nasdaq 100 ETF", "成长股敞口"], ["Quarterly Buffer ETFs", "10%至20% Buffer", "3个月", "S&P 500、Nasdaq 100、Russell 2000、MSCI EAFE、MSCI EM", "短周期系列"]], font_size=7.8)
    add_before(doc, ref, "资料来源：Innovator、Morningstar、易方达产品研究", "资料来源", size=9, color=GREY, after=6)

    add_heading_before(doc, ref, "2.4 NEOS：高收入叙事与差异化分派", 2)
    add_before(doc, ref, "NEOS在使用期权的衍生策略产品中AUM约248.74亿美元，样本内产品数量为8只。其代表产品包括SPYI和QQQI，产品定位更直接，围绕成熟宽基指数提供高收入策略，并把月度分派、税务处理和现金流属性作为投资者沟通重点。与J.P. Morgan主动权益组合加ELN的路径相比，NEOS更强调高收入叙事和期权收入的显性表达。")
    add_before(doc, ref, "表：NEOS高收入策略产品对比", "表格标题", size=10.5, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_doc_table(doc, ref, ["产品", "底层资产", "策略定位", "费率", "AUM", "需补来源"], [["SPYI", "S&P 500相关敞口", "高收入option income", "0.68%", "102.37亿美元", "NEOS官网"], ["QQQI", "Nasdaq 100相关敞口", "高收入option income", "0.68%", "129.01亿美元", "NEOS官网"]])
    add_before(doc, ref, "资料来源：Nasdaq、NEOS、Morningstar、易方达产品研究", "资料来源", size=9, color=GREY, after=6)

    add_heading_before(doc, ref, "2.5 Global X：规则化买写和风险管理收入矩阵", 2)
    add_before(doc, ref, "Global X在使用期权的衍生策略产品中AUM约134.06亿美元，样本内产品数量为18只。QYLD和XYLD是典型规则化buy-write产品线，分别围绕Nasdaq 100和S&P 500构建指数敞口并系统性卖出指数认购期权。与主动管理型JEPI和JEPQ相比，这类产品更强调透明、规则化和可复制。")
    add_before(doc, ref, "Global X还通过XRMI、QRMI和Tail Risk系列覆盖风险管理收入和保护性期权需求，说明成熟管理人通常不会只做单一covered call产品，而是围绕相同底层指数横向扩展收益增强、风险管理和尾部保护。Mirae Asset体系在香港市场推动备兑认购ETF本地化发行，可作为成熟策略向区域市场复制的案例，但按当前使用期权产品AUM口径并未进入全球前五。")
    add_before(doc, ref, "表：Global X备兑及风险管理策略产品对比", "表格标题", size=10.5, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_doc_table(doc, ref, ["产品", "管理人", "底层资产", "策略定位", "费率", "AUM"], [["QYLD", "Global X", "Nasdaq 100", "规则化buy-write", "0.60%", "83.16亿美元"], ["XYLD", "Global X", "S&P 500", "规则化buy-write", "0.60%", "31.82亿美元"], ["XRMI", "Global X", "S&P 500", "risk managed income", "0.60%", "0.50亿美元"], ["QRMI", "Global X", "Nasdaq 100", "risk managed income", "0.60%", "0.16亿美元"]])
    add_before(doc, ref, "资料来源：Nasdaq、Global X、Morningstar、易方达产品研究", "资料来源", size=9, color=GREY, after=6)

    doc.save(REPORT_PATH)


def update_markdown(detail):
    opt = detail[detail["是否使用期权"].eq("是")]
    non = detail[detail["是否使用期权"].eq("否")]
    total = detail["Fund Size USD"].sum()
    cross_count = pd.crosstab(detail["产品类型"], detail["是否使用期权"])
    cross_aum = pd.pivot_table(detail, index="产品类型", columns="是否使用期权", values="Fund Size USD", aggfunc="sum", fill_value=0)
    md = MD_PATH.read_text(encoding="utf-8") if MD_PATH.exists() else ""
    start = md.find("## 六、当前统计结果")
    if start != -1:
        md = md[:start]
    md += f"""## 六、当前统计结果

截至二级复核最终简表同步后，当前底稿统计为：

| 是否使用期权 | 产品数量 | AUM（亿美元） | AUM占比 |
|---|---:|---:|---:|
| 是 | {len(opt)} | {opt['Fund Size USD'].sum() / 1e8:.2f} | {opt['Fund Size USD'].sum() / total * 100:.2f}% |
| 否 | {len(non)} | {non['Fund Size USD'].sum() / 1e8:.2f} | {non['Fund Size USD'].sum() / total * 100:.2f}% |
| 合计 | {len(detail)} | {total / 1e8:.2f} | 100.00% |

按四大类交叉看：

| 产品类型 | 使用期权数量 | 未使用期权数量 | 使用期权AUM（亿美元） | 说明 |
|---|---:|---:|---:|---|
"""
    notes = {
        "收益增强型": "按当前规则全部视为使用期权",
        "风险缓冲型": "按当前规则全部视为使用期权",
        "杠杆与反向型": "绝大多数未使用期权，少数因明确期权线索或人工复核归为使用期权",
        "另类衍生策略型": "已完成二级复核，不再保留待人工核验项",
    }
    for typ in TYPE_ORDER:
        use_count = int(cross_count.loc[typ, "是"]) if "是" in cross_count.columns else 0
        no_count = int(cross_count.loc[typ, "否"]) if "否" in cross_count.columns else 0
        use_aum = float(cross_aum.loc[typ, "是"]) if "是" in cross_aum.columns else 0
        md += f"| {typ} | {use_count} | {no_count} | {use_aum / 1e8:.2f} | {notes[typ]} |\n"
    md += """
## 七、正文引用建议

正式报告中建议使用以下表述：

> 根据Morningstar字段规则识别，并结合灰区关键词复核及二级复核最终简表，当前衍生策略产品样本中，使用期权的产品共986只，AUM约3094.12亿美元；未使用期权的产品共2175只，AUM约3228.26亿美元。当前样本已无待人工核验产品。

不建议使用以下表述：

| 不建议表述 | 原因 |
|---|---|
| 已逐只完成官方核验 | 当前结果基于规则识别、重点灰区复核和二级复核最终简表，不代表3161只产品全部完成同等深度官方文件审阅 |
| 期权主导产品 | 当前口径为“是否使用期权”，并不判断期权在收益来源中的主导程度 |
| 未确认使用期权 | 容易混淆“否”和“待人工核验”；当前最终表同步后已不保留待人工核验项 |
"""
    MD_PATH.write_text(md, encoding="utf-8")


def main():
    BACKUP_DIR.mkdir(exist_ok=True)
    if not REPORT_BACKUP.exists():
        shutil.copy2(REPORT_PATH, REPORT_BACKUP)
    if not WORKBOOK_BACKUP.exists():
        shutil.copy2(WORKBOOK_PATH, WORKBOOK_BACKUP)
    detail = load_and_apply_final()
    stats = build_stats(detail)
    update_report(detail, stats)
    update_markdown(detail)
    print("updated")
    print(detail["是否使用期权"].value_counts().to_string())
    print(detail.groupby("是否使用期权")["Fund Size USD"].sum().div(1e8).to_string())


if __name__ == "__main__":
    main()
