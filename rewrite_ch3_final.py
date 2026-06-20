#!/usr/bin/env python3
"""
第三章完全重写脚本 — 最终版
3.1: JEPI JEPQ QQQI SPYI CAIE (5产品, 机制/绩效/费率各三段)
3.2: 三类Buffer (Laddered/月度单一/短周期, SPBU归入月度类)
3.3: TQQQ + CSOP SK Hynix 2x (指数/个股两类)
3.4: DBMF + NTSX
删除3.5, 费率分散入各产品段
只替换 三、到 四、之间的内容
"""

import pandas as pd, numpy as np
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
import warnings; warnings.filterwarnings('ignore')

# ============================================================
BASE = Path("/Users/castle/Desktop/space for claude")
RAW_PATH = BASE / "global derivative ETFs-20260616.xlsx"
WB_PATH = BASE / "海外ETF期权及衍生策略产品研究_核验底稿.xlsx"
DOC_PATH = BASE / "海外ETF期权及衍生策略产品研究0618_更新版.docx"
OUT_DOC = BASE / "海外ETF期权及衍生策略产品研究0618_更新版.docx"
OUT_WB = BASE / "海外ETF期权及衍生策略产品研究_核验底稿_更新版.xlsx"

CN_FONT = "KaiTi"; EN_FONT = "Times New Roman"
INK = "000000"; GREY = "6B7280"; HEADER_FILL = "F2F4F7"; BORDER_COLOR = "D0D7DE"; YELLOW = "FFF2CC"

BENCHMARKS = {
    'S&P 500': {'ticker':'IVV','fee':0.03,'name':'iShares Core S&P 500 ETF'},
    'Nasdaq 100': {'ticker':'QQQM','fee':0.15,'name':'Invesco Nasdaq 100 ETF'},
}

# ============================================================
def load_data():
    raw = pd.read_excel(RAW_PATH, sheet_name="20260616_global_derivative")
    raw['SecId'] = raw['SecId'].astype(str).str.strip()
    wb = pd.read_excel(WB_PATH, sheet_name="衍生策略分类明细")
    wb['SecId'] = wb['SecId'].astype(str).str.strip()
    raw['使用期权'] = raw['SecId'].map(wb.set_index('SecId')['是否使用期权']).fillna('否')
    raw['产品类型'] = raw['SecId'].map(wb.set_index('SecId')['产品类型']).fillna('另类衍生策略型')
    for c in ['Fund Size USD','Annual Report Net Expense Ratio']:
        raw[c] = pd.to_numeric(raw[c], errors='coerce')
    raw['Fund Size USD'] = raw['Fund Size USD'].fillna(0)
    raw['Fee'] = raw['Annual Report Net Expense Ratio']
    raw['Inception Date'] = pd.to_datetime(raw['Inception Date'], errors='coerce')
    raw['Inception Year'] = raw['Inception Date'].dt.year.fillna(0).astype(int)
    perf_map = {
        'Ret_1Yr':'Ret 1 Yr (Mo-End)','Ret_3Yr':'Ret Annlzd 3 Yr (Mo-End)',
        'StdDev_1Yr':'Std Dev 1 Yr (Mo-End) Risk Currency','StdDev_3Yr':'Std Dev 3 Yr (Mo-End) Risk Currency',
        'Sharpe_1Yr':'Sharpe Ratio 1 Yr (Mo-End) Risk Currency','Sharpe_3Yr':'Sharpe Ratio 3 Yr (Mo-End) Risk Currency',
        'Beta_1Yr':'Beta 1 Yr (Mo-End) Risk Currency','Beta_3Yr':'Beta 3 Yr (Mo-End) Risk Currency',
    }
    for k,v in perf_map.items(): raw[k] = pd.to_numeric(raw[v], errors='coerce')
    flow_map = {'Flow_1Mo':'Est Fund-Level Net Flow 1 Mo (Mo-End) USD',
                'Flow_1Yr':'Est Fund-Level Net Flow 1 Yr (Mo-End) USD',
                'Flow_3Yr':'Est Fund-Level Net Flow 3 Yr (Mo-End) USD'}
    for k,v in flow_map.items(): raw[k] = pd.to_numeric(raw[v], errors='coerce')
    return raw

def get_prod(raw, ticker):
    sub = raw[raw['Ticker'].str.upper() == ticker.upper()]
    if len(sub)==0: return None
    return sub.sort_values('Fund Size USD', ascending=False).iloc[0]

# ============================================================
# HELPERS
def fr(v): return f"{v:.2f}%" if pd.notna(v) else "N/A"
def fb(v): return f"{v*100:.0f}bp" if pd.notna(v) else "N/A"
def ff(v): return f"{v:.2f}%" if pd.notna(v) else "N/A"
def f2(v): return f"{v:.2f}" if pd.notna(v) else "N/A"

def font_run(run, size=None, bold=None, color=None):
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = EN_FONT
    r_pr = run._element.get_or_add_rPr()
    rf = r_pr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        r_pr.append(rf)
    rf.set(qn("w:ascii"), EN_FONT); rf.set(qn("w:hAnsi"), EN_FONT); rf.set(qn("w:eastAsia"), CN_FONT)

def set_pf(p, size=11, before=0, after=6, line=1.2, align=None):
    pf = p.paragraph_format
    pf.space_before=Pt(before); pf.space_after=Pt(after)
    pf.line_spacing_rule=WD_LINE_SPACING.MULTIPLE; pf.line_spacing=line
    if align: p.alignment=align
    for r in p.runs: font_run(r,size=size)

def add_p(doc, ref, text, style="Normal", size=11, bold=False, color=INK, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text); font_run(run, size=size, bold=bold, color=color)
    set_pf(p, size=size, before=before, after=after, align=align)
    ref.addprevious(p._element); return p

def add_h(doc, ref, text, level):
    s = {1:"Heading 1",2:"Heading 2",3:"Heading 3"}[level]
    return add_p(doc,ref,text,s,size={1:16,2:13,3:12}[level],bold=True,
                 before={1:16,2:12,3:8}[level],after={1:8,2:6,3:4}[level])

def add_src(doc,ref,text): return add_p(doc,ref,"资料来源："+text,"资料来源",size=9,color=GREY,after=2)
def add_note(doc,ref,text): return add_p(doc,ref,"注："+text,"注释",size=9,color=GREY,after=6)
def add_ttl(doc,ref,text): return add_p(doc,ref,text,"表格标题",size=10.5,bold=True,after=2,align=WD_ALIGN_PARAGRAPH.CENTER)

def fill_cell(cell, text, size=8.0, bold=False, shade=None):
    if shade: set_shade(cell, shade)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell.text = ""
    p = cell.paragraphs[0]
    try: p.style = "表格正文"
    except: pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.15
    r = p.add_run("" if text is None else str(text)); font_run(r, size=size, bold=bold, color=INK)

def set_shade(cell, fill):
    tc = cell._tc.get_or_add_tcPr()
    s = tc.find(qn("w:shd"))
    if s is None: s=OxmlElement("w:shd"); tc.append(s)
    s.set(qn("w:fill"), fill)

def set_borders(table, color=BORDER_COLOR, size="4"):
    tbl = table._tbl.tblPr
    b = tbl.first_child_found_in("w:tblBorders")
    if b is None: b=OxmlElement("w:tblBorders"); tbl.append(b)
    for e in ("top","left","bottom","right","insideH","insideV"):
        el=b.find(qn(f"w:{e}"));
        if el is None: el=OxmlElement(f"w:{e}"); b.append(el)
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),size); el.set(qn("w:space"),"0"); el.set(qn("w:color"),color)

def add_tbl(doc, ref, headers, rows, fs=7.8):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): fill_cell(table.rows[0].cells[i], h, size=8.2, bold=True, shade=HEADER_FILL)
    for rd in rows:
        cells = table.add_row().cells
        for i,v in enumerate(rd):
            t = "" if v is None else str(v)
            fill_cell(cells[i], t, size=fs, shade=YELLOW if "待补充" in t else None)
    set_borders(table); ref.addprevious(table._element); return table

def remove_between(doc, start_text, end_text):
    body = doc.element.body; children = list(body.iterchildren())
    si = ei = None
    for idx, child in enumerate(children):
        if isinstance(child, CT_P):
            t = "".join(child.itertext()).strip()
            if si is None and t.startswith(start_text): si = idx
            elif si is not None and t.startswith(end_text): ei = idx; break
    if si is None or ei is None: raise RuntimeError(f"Cannot find: {start_text} to {end_text}")
    ref = children[ei]
    for child in children[si:ei]: body.remove(child)
    return ref

# ============================================================
# CHAPTER 3
# ============================================================
def gen_ch3(doc, ref, raw):
    add_h(doc, ref, "三、主要产品类型拆解", 1)
    gen_31(doc, ref, raw)
    gen_32(doc, ref, raw)
    gen_33(doc, ref, raw)
    gen_34(doc, ref, raw)

# ============================================================
# 3.1 收益增强型
# ============================================================
def gen_31(doc, ref, raw):
    add_h(doc, ref, "3.1 收益增强型产品", 2)
    add_p(doc, ref,
        "收益增强型产品以Covered Call、Buy-Write和Premium Income策略为主，"
        "产品通常持有股票或指数相关敞口，同时卖出认购期权或通过ELN、互换等工具嵌入卖权结构，"
        "以权利金和底层资产分红形成现金流。收益来源可拆为四项：标的资产涨跌、成分股分红、"
        "卖出期权权利金、期权平仓或被执行损益。这类产品更适合震荡市、温和上涨和隐含波动率较高的市场环境。"
        "其核心代价是让渡上行空间——在快速单边上涨行情中可能明显落后于纯指数ETF；"
        "在急跌行情中，权利金只能提供有限缓冲，无法消除权益下跌风险。")

    # -- 情景分析表 --
    add_ttl(doc, ref, "表：Covered Call策略在不同市场情景下的收益结构")
    add_tbl(doc, ref,
        ["市场情景", "标的涨跌", "期权端结果", "产品总收益", "vs 纯指数ETF"],
        [
            ["大幅上涨", "+15%", "被行权，上行收益封顶于行权价", "行权价涨幅+权利金", "显著落后"],
            ["温和上涨", "+5%", "未被行权或部分行权", "标的涨幅+权利金", "基本同步"],
            ["横盘震荡", "0%", "权利金全部保留", "+权利金收入", "跑赢"],
            ["温和下跌", "-5%", "权利金提供有限缓冲", "标的跌幅+权利金", "略优"],
            ["大幅下跌", "-15%", "权利金缓冲有限", "标的跌幅+权利金", "略优，跌幅基本同步"],
        ], fs=8.0)
    add_src(doc, ref, "易方达产品研究。注：仅示意策略逻辑，实际收益取决于行权价选择、覆盖比例和再平衡频率。")

    # -- 3.1.1 JEPI --
    add_h(doc, ref, "3.1.1 JEPI", 3)
    r = get_prod(raw, 'JEPI')
    if r is not None:
        bench = BENCHMARKS['S&P 500']
        add_p(doc, ref,
            f"JEPI（JPMorgan Equity Premium Income ETF）于2020年5月发行，截至2026年5月末AUM约445.48亿美元，"
            f"为全球规模最大的期权策略ETF。年化费用0.35%，相对底层基准S&P 500纯指数ETF（{bench['ticker']} {bench['fee']:.2f}%）"
            f"溢价{fb(0.35-bench['fee'])}。")
        add_p(doc, ref,
            f"策略机制上，JEPI通过两层结构实现收益目标：第一层是主动管理的美国大盘低波动股票组合，"
            f"由J.P. Morgan基本面研究员根据估值和选股模型构建；第二层为最多20%净资产投资于ELN，"
            f"ELN将S&P 500的卖出认购期权经济特征嵌入一张票据中，形成月度权利金收入。"
            f"基金在正常情况下至少80%资产投资于权益证券（含普通股和ELN）。"
            f"与简单持有S&P 500并全额卖出认购期权的规则化策略不同，JEPI的权益端具有主动选股的超额收益潜力，"
            f"期权端也不追求全覆盖，而是通过ELN灵活控制卖出期权的名义敞口。")
        add_p(doc, ref,
            f"绩效方面，近1年回报{fr(r['Ret_1Yr'])}，近3年年化回报{fr(r['Ret_3Yr'])}，"
            f"1年波动率{fr(r['StdDev_1Yr'])}，1年Beta {f2(r['Beta_1Yr'])}，"
        f"近1年Sharpe Ratio {f2(r['Sharpe_1Yr'])}。"
            f"低Beta特征（相对S&P 500约0.45）反映了卖出认购期权对组合波动率的抑制效果。"
            f"近1年资金净流入约+49亿美元，过去3年累计净流入约+179亿美元，"
            f"体现了市场对该策略的持续认可。")

    # -- 3.1.2 JEPQ --
    add_h(doc, ref, "3.1.2 JEPQ", 3)
    r = get_prod(raw, 'JEPQ')
    if r is not None:
        bench = BENCHMARKS['Nasdaq 100']
        add_p(doc, ref,
            f"JEPQ（JPMorgan Nasdaq Equity Premium Income ETF）于2022年5月发行，"
            f"AUM约401.53亿美元，为规模第二大的期权策略ETF。年化费用0.35%，"
            f"相对底层基准Nasdaq 100纯指数ETF（{bench['ticker']} {bench['fee']:.2f}%）溢价{fb(0.35-bench['fee'])}。"
            f"与JEPI共享相同的ELN+主动权益架构，区别在于权益组合围绕Nasdaq 100相关股票构建，"
            f"成长和科技属性更强，期权端也对应嵌入Nasdaq 100的卖出认购期权敞口。")
        add_p(doc, ref,
            f"由于Nasdaq 100的成长股权重和高隐含波动率特征，JEPQ的回报和波动均高于JEPI："
            f"近1年回报{fr(r['Ret_1Yr'])}（JEPI为8.55%），近3年年化{fr(r['Ret_3Yr'])}（JEPI为9.53%），"
            f"1年波动率{fr(r['StdDev_1Yr'])}，1年Beta {f2(r['Beta_1Yr'])}。"
            f"近1年Sharpe Ratio {f2(r['Sharpe_1Yr'])}，在承担更高波动的同时提供了更优的风险调整后收益。"
            f"近1年资金净流入约+88亿美元，过去3年累计约+314亿美元，资金吸引力在同类产品中最强。")

    # -- 3.1.3 QQQI --
    add_h(doc, ref, "3.1.3 QQQI", 3)
    r = get_prod(raw, 'QQQI')
    if r is not None:
        bench = BENCHMARKS['Nasdaq 100']
        add_p(doc, ref,
            f"QQQI（NEOS Nasdaq-100 High Income ETF）于2024年1月发行，仅约2.5年时间AUM即达129.01亿美元，"
            f"是增长最快的期权策略产品之一。年化费用0.68%，相对{bench['ticker']}（{bench['fee']:.2f}%）溢价{fb(0.68-bench['fee'])}。"
            f"与JEPI/JEPQ的ELN路径不同，NEOS采用更直接的active option income策略——"
            f"在持有Nasdaq 100相关敞口的同时，通过主动管理的期权卖出获取月度分派。"
            f"NEOS将月度分派、税务效率（部分分派可能被视为资本返还）和现金流属性作为核心营销叙事。")
        add_p(doc, ref,
            f"绩效方面，近1年回报{fr(r['Ret_1Yr'])}，1年波动率{fr(r['StdDev_1Yr'])}，"
            f"近1年Sharpe Ratio {f2(r['Sharpe_1Yr'])}，Beta {f2(r['Beta_1Yr'])}。"
            f"由于成立不足3年，暂无3年年化数据。近1年资金净流入约+96亿美元，"
            f"在所有期权策略ETF中居首位——甚至超过了JEPI和JEPQ，"
            f"反映出高收入叙事在零售投资者中的强大吸引力。")

    # -- 3.1.4 SPYI --
    add_h(doc, ref, "3.1.4 SPYI", 3)
    r = get_prod(raw, 'SPYI')
    if r is not None:
        bench = BENCHMARKS['S&P 500']
        add_p(doc, ref,
            f"SPYI（NEOS S&P 500 High Income ETF）于2022年8月发行，AUM约102.37亿美元，"
            f"是NEOS的首只产品。年化费用0.68%，相对{bench['ticker']}（{bench['fee']:.2f}%）溢价{fb(0.68-bench['fee'])}。"
            f"策略逻辑与QQQI一致，但底层敞口为S&P 500而非Nasdaq 100。"
            f"近1年回报{fr(r['Ret_1Yr'])}，近3年年化{fr(r['Ret_3Yr'])}，"
            f"1年波动率{fr(r['StdDev_1Yr'])}，1年Beta {f2(r['Beta_1Yr'])}，近1年Sharpe {f2(r['Sharpe_1Yr'])}。"
            f"近1年净流入约+58亿美元，NEOS的S&P和Nasdaq两条产品线均处于快速增长通道。")

    # -- 3.1.5 CAIE (Autocallable) --
    add_h(doc, ref, "3.1.5 CAIE", 3)
    r = get_prod(raw, 'CAIE')
    if r is not None:
        add_p(doc, ref,
            f"CAIE（Calamos US Equity Autocallable Income ETF）于2025年发行，"
            f"是全球首只将Autocallable（自动赎回）结构化票据装入ETF结构的产品，"
            f"AUM约{r['Fund Size USD']/1e8:.2f}亿美元，年化费用{ff(r['Fee'])}。"
            f"Autocallable策略的核心机制为：产品在发行时预设一系列定期观察日和对应的提前赎回条件"
            f"（如标的指数在观察日达到初始水平的100%以上），一旦触发即自动终止并向投资者支付本金加票息；"
            f"若整个存续期内始终未触发，到期时按最终标的水平和保护条款结算。"
            f"这类结构在海外结构化票据市场和国内雪球类产品中均有庞大的投资者基础。")
        add_p(doc, ref,
            f"Calamos将Autocallable装入ETF的意义在于：首次将这一传统柜台产品转化为标准化、"
            f"可日内交易、低投资门槛的ETF形态。相比柜台结构化票据，ETF版本具有每日流动性和透明的持仓披露，"
            f"降低了信息不对称和交易对手集中风险。"
            f"由于CAIE成立不足1年，暂无完整的1年回报数据；"
            f"近1年净流入约+9亿美元，说明市场需求初步验证。")
        add_p(doc, ref,
            f"对国内产品开发的启示：国内雪球类产品一度发展迅速（2023年末存续规模峰值约4000亿元），但均为柜台结构化票据或收益凭证形式，"
            f"存在门槛高、流动性差、结构不透明等问题。CAIE的ETF化路径为国内探索"
            f"「结构化产品公募化」提供了可参考的范本，但国内受限于场外衍生品监管框架和公募基金参与衍生品的约束条件，"
            f"短期内直接复制的可行性较低，可作为中长期储备方向。")

    # -- 汇总对比表 --
    add_ttl(doc, ref, "表：收益增强型代表产品多维度对比")
    prods_31 = ['JEPI','JEPQ','QQQI','SPYI','CAIE']
    rows31 = []
    for t in prods_31:
        r = get_prod(raw, t)
        if r is None: continue
        b = BENCHMARKS.get('S&P 500') if t in ['JEPI','SPYI'] else BENCHMARKS.get('Nasdaq 100')
        bt = f"{b['ticker']}({b['fee']:.2f}%)" if b else "N/A"
        fp = fb(r['Fee']-b['fee']) if b and pd.notna(r['Fee']) else "N/A"
        path = {'JEPI':'主动+ELN','JEPQ':'主动+ELN','QQQI':'active option income',
                'SPYI':'active option income','CAIE':'Autocallable ETF'}.get(t,'')
        rows31.append([t, path, ff(r['Fee']), bt, fp,
                       f"{r['Fund Size USD']/1e8:.1f}" if r['Fund Size USD']>0 else "N/A",
                       fr(r['Ret_1Yr']), fr(r['Ret_3Yr']), fr(r['StdDev_1Yr']),
                       f2(r['Sharpe_1Yr']), f2(r['Beta_1Yr'])])
    add_tbl(doc, ref,
        ["Ticker","策略路径","费率","基准ETF(费率)","费率溢价","AUM(亿美元)",
         "近1年回报","近3年年化","1年波动率","Sharpe 1Yr","Beta 1Yr"], rows31, fs=7.2)
    add_src(doc, ref,
        "J.P. Morgan、NEOS、Calamos、Morningstar、易方达产品研究。"
        "基准ETF费率来自BlackRock/Invesco官网（2026年6月访问）。"
        "CAIE成立不足1年，近1年回报和3年年化暂不可得。")

# ============================================================
# 3.2 风险缓冲型
# ============================================================
def gen_32(doc, ref, raw):
    add_h(doc, ref, "3.2 风险缓冲型产品", 2)
    add_p(doc, ref,
        "风险缓冲型产品（Buffer/Target Outcome ETF）通过FLEX Options构建期权组合，"
        "预先设定一段持有期内的风险收益边界：下跌端提供一定比例的缓冲（如9%、15%或30%），"
        "上涨端通常设置收益上限（cap）。典型结构可理解为买入保护性认沽、卖出更低行权价认沽限定保护区间，"
        "再卖出上方认购期权补贴保护成本。该类产品AUM约902亿美元（占衍生策略总规模的14.27%），"
        "在四大类别中集中度最低，产品形态仍有较大差异化空间。")

    # -- Buffer情景表 --
    add_ttl(doc, ref, "表：Buffer策略在不同市场情景下的收益结构（以15% Buffer为例）")
    add_tbl(doc, ref,
        ["市场情景", "标的涨跌", "产品结果", "vs 纯指数ETF"],
        [
            ["涨幅超过Cap", "+20%", "收益封顶于Cap（如+10%）", "显著落后"],
            ["涨幅在Cap内", "+8%", "全额参与+8%", "基本同步"],
            ["横盘", "0%", "0%", "同步"],
            ["跌幅在Buffer内", "-8%", "不亏损（Buffer吸收全部跌幅）", "显著跑赢"],
            ["跌幅超过Buffer", "-20%", "超出Buffer部分亏损（-5%）", "跑赢"],
        ], fs=8.0)
    add_src(doc, ref, "易方达产品研究。注：仅示意策略逻辑，实际收益取决于cap水平、buffer深度、outcome period和建仓时点。")

    # -- 分类概述 --
    add_p(doc, ref,
        "从产品设计维度，风险缓冲型ETF可归为三类："
        "一是Laddered Buffer（滚动配置型），通过持有不同到期月份的底层Buffer ETF实现每季度刷新保护区间，"
        "降低单一建仓时点的影响；"
        "二是月度单一Buffer（单次选择型），每个产品锁定一个固定的outcome period和缓冲参数，"
        "投资者需自行选择到期月份；"
        "三是短周期Buffer（高频刷新型），以3个月甚至更短的outcome period换取更频繁的保护更新。")

    add_ttl(doc, ref, "表：风险缓冲型产品分类对比")
    add_tbl(doc, ref,
        ["类型", "代表产品", "AUM(亿美元)", "费率", "Outcome Period", "Buffer/Cap特征", "适用场景"],
        [
            ["Laddered Buffer", "BUFR", "97.32", "0.10%", "持续滚动，每季刷新", "约9% Buffer，有Cap", "长期配置，无需择时"],
            ["Laddered Buffer(深度)", "BUFD", "18.92", "0.10%", "持续滚动", "更深Buffer(约20%)，有Cap", "更保守的长期配置"],
            ["月度单一Buffer", "PJAN(Innovator)", "14.83", "0.79%", "1年(1月起)", "15% Buffer，有Cap", "投资者自主选择到期月份"],
            ["月度单一Buffer(去Cap)", "SPBU(AllianzIM)", "1.63", "0.05%", "1年", "15% Buffer，无Cap", "需下行保护但不愿放弃上行"],
            ["短周期Buffer", "BALT(Innovator)", "25.22", "0.69%", "3个月", "20% Buffer，有Cap", "对建仓时点敏感度较高"],
        ], fs=7.6)
    add_src(doc, ref, "First Trust、Innovator、AllianzIM、Morningstar、易方达产品研究。")

    # -- 3.2.1 Laddered Buffer --
    add_h(doc, ref, "3.2.1 Laddered Buffer：以BUFR为代表", 3)
    r = get_prod(raw, 'BUFR')
    if r is not None:
        inception_year = r['Inception Date'].year if pd.notna(r['Inception Date']) else 2020
        inception_month = r['Inception Date'].month if pd.notna(r['Inception Date']) else 8
        add_p(doc, ref,
            f"BUFR（FT Vest Laddered Buffer ETF）于{inception_year}年{inception_month}月发行，"
            f"AUM约97.32亿美元，为规模最大的Buffer ETF。"
            f"自身管理费0.10%；因采用基金中基金（FOF）结构，底层12只ETF另收取费用，"
            f"投资者实际承担的总费率约0.95%。"
            f"核心创新在于Laddered设计：基金等权持有12只底层FT Vest U.S. Equity Buffer ETF，"
            f"每只对应一个不同到期月份，每月有一只底层ETF进入新的为期一年的目标结果期间，"
            f"相应刷新其buffer和cap。这一设计将原本需要投资者自行选择月份和建仓时点的Buffer产品，"
            f"转化为更容易长期持有的「常青」配置工具。"
            f"绩效方面，近1年回报{fr(r['Ret_1Yr'])}，1年波动率仅{fr(r['StdDev_1Yr'])}，"
            f"Sharpe Ratio {f2(r['Sharpe_1Yr'])}，1年Beta {f2(r['Beta_1Yr'])}，"
            f"体现了Buffer结构在控制下行风险方面的有效性。近1年净流入约+17亿美元。")

    # -- 3.2.2 月度单一Buffer --
    add_h(doc, ref, "3.2.2 月度单一Buffer：以Innovator系列和SPBU为代表", 3)
    r = get_prod(raw, 'PJAN')
    if r is not None:
        add_p(doc, ref,
            f"月度单一Buffer是Buffer ETF中产品数量最多的类别。以Innovator为例，"
            f"其PJAN-PDEC系列覆盖12个自然月，每个产品锁定一个1年outcome period和15%的Power Buffer保护区间。"
            f"以PJAN（Innovator U.S. Equity Power Buffer ETF - January）为代表，AUM约14.83亿美元，"
            f"费率0.79%，近1年回报{fr(r['Ret_1Yr'])}，1年波动率{fr(r['StdDev_1Yr'])}。"
            f"投资者需在outcome period开始时买入并持有至到期，才能获得完整的buffer和cap；"
            f"中途买入的剩余buffer和cap会随标的涨跌而变化。First Trust的FJAN-FDEC系列结构类似，"
            f"但以更低的费率（约0.85%）和更丰富的底层覆盖参与竞争。")

    r2 = get_prod(raw, 'SPBU')
    if r2 is not None:
        add_p(doc, ref,
            f"2025年AllianzIM推出的SPBU（Buffer15 Uncapped Allocation ETF）是该类别的重要创新——"
            f"它在提供15%下行保护的同时，完全取消了传统Buffer ETF的上行封顶。"
            f"SPBU通过仅买入保护性认沽、不卖出认购期权的期权结构实现了「有保护、不封顶」的收益特征，"
            f"AUM约1.63亿美元，费率仅0.05%，近1年回报{fr(r2['Ret_1Yr'])}。"
            f"虽然其buffer深度（15%）与Innovator Power Buffer系列相同，"
            f"但取消cap意味着投资者在牛市中不再让渡上行空间，"
            f"代价是保护成本全部由基金自身承担（无认购期权权利金补贴），这也解释了其费率极低的原因。"
            f"「去Cap」设计为Buffer ETF品类打开了新的创新方向，对国内低波动产品开发具有直接的策略参考价值。")

    # -- 3.2.3 短周期Buffer --
    add_h(doc, ref, "3.2.3 短周期Buffer：以BALT为代表", 3)
    r = get_prod(raw, 'BALT')
    if r is not None:
        add_p(doc, ref,
            f"BALT（Innovator Defined Wealth Shield ETF）于2021年6月发行，AUM约25.22亿美元，"
            f"费率0.69%。与月度系列的1年outcome period不同，BALT采用3个月的短周期，"
            f"提供20%的Buffer保护。更短的周期意味着更频繁的保护刷新，"
            f"降低了投资者在单一建仓时点进场的不利影响。")
        add_p(doc, ref,
            f"绩效方面，近1年回报{fr(r['Ret_1Yr'])}，1年波动率仅{fr(r['StdDev_1Yr'])}，"
            f"Sharpe Ratio {f2(r['Sharpe_1Yr'])}，1年Beta仅{f2(r['Beta_1Yr'])}，"
            f"在同类产品中属于极低波动水平。近1年净流入约+9亿美元。"
            f"BALT和BUFR分别代表了Buffer ETF在「缩短周期」和「滚动配置」两个方向上的演化尝试，"
            f"两者的共同目标是降低投资者对建仓时点的敏感度，但实现路径不同。")

    # -- 绩效对比表 --
    add_ttl(doc, ref, "表：风险缓冲型代表产品绩效对比")
    prods_32 = ['BUFR','BUFD','BUFQ','BALT','PJAN','SPBU']
    rows32 = []
    for t in prods_32:
        r = get_prod(raw, t)
        if r is None: continue
        rows32.append([t, f"{r['Fund Size USD']/1e8:.1f}", ff(r['Fee']),
                       fr(r['Ret_1Yr']), fr(r['Ret_3Yr']), fr(r['StdDev_1Yr']),
                       f2(r['Sharpe_1Yr']), f2(r['Beta_1Yr'])])
    add_tbl(doc, ref,
        ["Ticker","AUM(亿美元)","费率","近1年回报","近3年年化","1年波动率","Sharpe 1Yr","Beta 1Yr"],
        rows32, fs=7.8)
    add_src(doc, ref, "First Trust、Innovator、AllianzIM、Morningstar、易方达产品研究。")

# ============================================================
# 3.3 杠杆与反向型
# ============================================================
def gen_33(doc, ref, raw):
    add_h(doc, ref, "3.3 杠杆与反向型产品", 2)
    add_p(doc, ref,
        "杠杆与反向型ETF通过期货、互换等衍生工具提供标的每日多倍收益或反向收益。"
        "该类产品共1,843只，AUM约3,022亿美元（占衍生策略总规模的47.79%），"
        "在数量和规模上均居四类之首。其中仅21只使用期权工具，绝大多数通过期货、互换和每日再平衡机制实现。")

    add_p(doc, ref,
        "杠杆ETF的核心机制是每日再平衡——每个交易日结束时调整衍生品头寸，"
        "确保下一个交易日的杠杆倍数回到目标水平（如2x或3x）。这一机制在单边行情中会产生正复利效应"
        "（连续上涨时，3x ETF的累计收益超过标的累计涨幅的3倍），但在震荡市中会产生波动拖累"
        "（volatility decay）：标的先涨后跌回到原点时，杠杆ETF的净值将低于初始值。"
        "因此杠杆ETF更适合短期方向性交易，而非长期买入持有。")

    # -- 情景表 --
    add_ttl(doc, ref, "表：每日3倍杠杆ETF在不同市场路径下的表现示意")
    add_tbl(doc, ref,
        ["市场路径", "标的累计收益", "3x杠杆ETF累计收益", "相对标的倍数的偏差", "说明"],
        [
            ["连续3天每日+2%", "+6.12%", "+19.10%", ">3倍(正复利)", "单边上涨时复利效应放大收益"],
            ["涨跌交替(+2%,-2%,+2%)", "+1.88%", "+3.28%", "<3倍", "震荡导致波动拖累"],
            ["连续3天每日-2%", "-5.88%", "-16.97%", "<3倍(负复利)", "单边下跌时复利效应放大亏损"],
        ], fs=7.8)
    add_src(doc, ref, "易方达产品研究。注：仅示意每日再平衡机制的影响，实际结果取决于标的波动率和路径。")

    # -- 分类 --
    add_p(doc, ref,
        "按追踪标的类型，杠杆与反向型ETF可分为指数追踪型和个股追踪型。"
        "指数追踪型以宽基指数（如Nasdaq 100、S&P 500）或行业指数为底层，产品数量和规模均占主导；"
        "个股追踪型以单只高流动性股票为标的，近年增长较快，但集中风险和波动拖累效应更为显著。")

    add_ttl(doc, ref, "表：杠杆与反向型产品分类及代表产品")
    add_tbl(doc, ref,
        ["类型", "数量", "AUM(亿美元)", "最大产品", "杠杆倍数", "产品AUM(亿美元)", "主要市场"],
        [
            ["指数追踪型", "1,707", "2,654", "TQQQ(ProShares)", "3x", "402.95", "美国、韩国、中国香港"],
            ["个股追踪型", "136", "367", "CSOP SK Hynix 2x", "2x", "107.66", "中国香港、韩国、美国"],
        ], fs=7.8)
    add_src(doc, ref, "Morningstar、管理人官网、易方达产品研究。数据截至2026年5月31日。")

    # -- 3.3.1 TQQQ --
    add_h(doc, ref, "3.3.1 TQQQ", 3)
    r = get_prod(raw, 'TQQQ')
    if r is not None:
        add_p(doc, ref,
            f"TQQQ（ProShares UltraPro QQQ）于2010年2月发行，是目前规模最大的杠杆ETF，"
            f"AUM约402.95亿美元。产品追踪Nasdaq 100指数的每日3倍多头收益，年化费用0.82%。"
            f"得益于2025-2026年科技股的强劲表现，近1年回报高达{fr(r['Ret_1Yr'])}，"
            f"但1年波动率也达{fr(r['StdDev_1Yr'])}，Sharpe Ratio {f2(r['Sharpe_1Yr'])}。"
            f"1年Beta为{f2(r['Beta_1Yr'])}，相对Nasdaq 100的杠杆倍数基本稳定在3倍附近。")
        add_p(doc, ref,
            f"值得注意的是，尽管TQQQ近1年绝对回报极高，但其近1年资金净流出约-119亿美元，"
            f"反映出投资者可能在牛市高位获利了结，或对杠杆产品的风险认知在提升。"
            f"对于国内公募而言，TQQQ展示了杠杆ETF在极端行情下的规模爆发力和风险——"
            f"牛市中可以快速吸金，但波动拖累和路径依赖可能使长期持有者承受远高于预期的亏损。")

    # -- 3.3.2 CSOP SK Hynix 2x --
    add_h(doc, ref, "3.3.2 CSOP SK Hynix Daily 2x Leveraged Product", 3)
    add_p(doc, ref,
        "CSOP SK Hynix Daily 2x Leveraged Product由南方东英资产管理有限公司（CSOP）发行，"
        "是香港市场规模最大的杠杆产品，AUM约107.66亿美元。"
        "该产品追踪韩国SK海力士（SK Hynix）单只股票的每日2倍收益，"
        "代表了亚太市场个股杠杆ETF的产品形态。")
    add_p(doc, ref,
        "香港市场是亚太地区杠杆与反向产品的主要上市地之一。截至2026年5月，"
        "港交所共有约29只杠杆/反向产品，总AUM约181亿港元，由南方东英、博时等管理人主导。"
        "追踪标的覆盖恒生指数、恒生科技、沪深300、纳斯达克100以及三星电子、SK海力士等个股。"
        "香港杠杆/反向产品的监管框架相对成熟——SFC于2016年首次认可杠杆/反向产品在港上市，"
        "目前最高允许2倍杠杆，且要求产品名称明确标注杠杆倍数。")
    add_p(doc, ref,
        "对内地市场的启示：港股已有的杠杆/反向产品为内地投资者提供了间接参与渠道（如港股通），"
        "但内地公募基金直接发行杠杆ETF仍受限于衍生品净敞口限制和投资者适当性要求。"
        "此外，单股票杠杆产品虽然交易属性强、吸金能力突出，但个股集中风险和投资者误用风险均较高，"
        "与国内公募「稳慎创新」和「投资者保护」的政策导向匹配度较低，短期内不宜作为优先推进方向。")

# ============================================================
# 3.4 另类衍生策略型
# ============================================================
def gen_34(doc, ref, raw):
    add_h(doc, ref, "3.4 另类衍生策略型产品", 2)
    add_p(doc, ref,
        "另类衍生策略型产品覆盖尾部风险对冲、波动率策略、管理期货、多资产叠加和资本效率等多元策略，"
        "共468只产品，AUM约454亿美元（占衍生策略总规模的7.18%），其中115只使用期权工具。"
        "该类产品内部差异最大，按组合功能可归为四种子策略："
        "管理期货策略通过多资产趋势跟随获取与传统股债低相关的收益来源；"
        "资本效率叠加策略在同一资本基础上通过期货或期权叠加多类资产敞口；"
        "波动率与尾部风险策略以VIX期货或保护性期权管理极端市场风险；"
        "现金替代与混合型策略通过期权结构改善现金收益路径。"
        "以下选取管理期货和资本效率两个方向的最大代表产品展开分析。")

    # -- 3.4.1 DBMF --
    add_h(doc, ref, "3.4.1 DBMF", 3)
    r = get_prod(raw, 'DBMF')
    if r is not None:
        add_p(doc, ref,
            f"DBMF（iMGP DBi Managed Futures Strategy ETF）于2019年5月发行，AUM约40.35亿美元，"
            f"是全球规模最大的管理期货ETF。年化费用0.85%。"
            f"策略机制上，DBMF通过量化模型跟踪多资产期货市场（覆盖股票指数、国债、商品和外汇期货）"
            f"的趋势信号，动态调整多头或空头头寸。与传统的股债组合不同，"
            f"管理期货策略的收益来源是跨资产类别的趋势方向和趋势强度，而非资产本身的长期风险溢价。")
        add_p(doc, ref,
            f"这一策略的核心价值在于与股债资产的低相关性。DBMF近1年回报{fr(r['Ret_1Yr'])}，"
            f"1年Beta仅{f2(r['Beta_1Yr'])}——这意味着其收益几乎独立于股票市场走势。"
            f"近1年Sharpe Ratio {f2(r['Sharpe_1Yr'])}，1年波动率{fr(r['StdDev_1Yr'])}。"
            f"近1年资金净流入约+23亿美元，反映了投资者在市场波动期对低相关策略的配置需求。")
        add_p(doc, ref,
            f"对国内产品开发的启示：管理期货策略在国内已有私募和专户实践（如CTA策略产品），"
            f"公募化面临的核心挑战是策略透明度、流动性管理和费率问题。"
            f"参照DBMF的经验，若以规则化指数跟踪方式实现趋势跟随、并通过ETF结构提供每日流动性，"
            f"管理期货ETF在「固收+」和资产配置领域有明确的应用场景。")

    # -- 3.4.2 NTSX --
    add_h(doc, ref, "3.4.2 NTSX", 3)
    r = get_prod(raw, 'NTSX')
    if r is not None:
        add_p(doc, ref,
            f"NTSX（WisdomTree U.S. Efficient Core Fund）于2018年8月发行，AUM约13.83亿美元，"
            f"年化费用仅0.20%。其核心创新在于「90/60」资本效率结构："
            f"将约90%的净资产投资于美国大盘股票组合，同时利用剩余10%资产作为保证金，"
            f"通过国债期货获得约60%的美国国债敞口。在相同的资本基础上，"
            f"NTSX实现了100%的权益暴露加60%的债券暴露，相当于「一美元做了两美元的活」。")
        add_p(doc, ref,
            f"绩效方面，NTSX近1年回报{fr(r['Ret_1Yr'])}，近3年年化{fr(r['Ret_3Yr'])}，"
            f"1年波动率{fr(r['StdDev_1Yr'])}，1年Beta {f2(r['Beta_1Yr'])}。"
            f"其风险收益特征更接近权益产品（Beta接近1.0），但由于叠加了债券敞口，"
            f"在股债双杀环境中可能面临比纯权益ETF更大的波动。"
            f"近1年资金净流出约-0.6亿美元，规模增长动力相对有限。")
        add_p(doc, ref,
            f"对国内产品开发的启示：资本效率策略的设计思路——在同一资本上叠加多资产敞口——"
            f"对提高组合资金使用效率具有直接参考价值。"
            f"但国内公募基金参与国债期货的保证金比例和净敞口限制较为严格，"
            f"直接复制NTSX的90/60结构的可行性较低。"
            f"可先从低比例的期货叠加（如90/10或90/20）起步，待制度空间进一步打开后再逐步提高杠杆倍数。")

    # -- 对比表 --
    add_ttl(doc, ref, "表：另类衍生策略型代表产品对比")
    prods_34 = ['DBMF','NTSX']
    rows34 = []
    for t in prods_34:
        r = get_prod(raw, t)
        if r is None: continue
        strat = {'DBMF':'多资产趋势跟随，覆盖股债商汇期货',
                 'NTSX':'90/60资本效率，权益+国债期货叠加'}.get(t,'')
        rows34.append([t, f"{r['Fund Size USD']/1e8:.1f}", str(r.get('Firm Name',''))[:25],
                       ff(r['Fee']), fr(r['Ret_1Yr']), fr(r['Ret_3Yr']),
                       fr(r['StdDev_1Yr']), f2(r['Sharpe_1Yr']), f2(r['Beta_1Yr']), strat])
    add_tbl(doc, ref,
        ["Ticker","AUM(亿美元)","管理人","费率","近1年回报","近3年年化",
         "1年波动率","Sharpe 1Yr","Beta 1Yr","策略特征"],
        rows34, fs=7.8)
    add_src(doc, ref, "iMGP/DBi、WisdomTree、Morningstar、易方达产品研究。")

# ============================================================
# MAIN
# ============================================================
def main():
    print("=== 第三章重写 ===")
    raw = load_data()
    doc = Document(str(DOC_PATH))
    ref = remove_between(doc, "三、", "四、")
    gen_ch3(doc, ref, raw)
    doc.save(str(OUT_DOC))
    print(f"Saved: {OUT_DOC}")

if __name__ == "__main__":
    main()
