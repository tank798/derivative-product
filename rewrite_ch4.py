#!/usr/bin/env python3
"""
第四章重写脚本 — 增强版
结构: 开头(结论先行) → 4.1政策 → 4.2为什么没做起来 → 4.3落地建议
删除原第五章，原第六章改为第五章（由外部脚本处理）
"""
import pandas as pd, numpy as np
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor
import warnings; warnings.filterwarnings('ignore')

BASE = Path("/Users/castle/Desktop/space for claude")
DOC_PATH = BASE / "海外ETF期权及衍生策略产品研究0618_更新版.docx"
OUT_DOC = BASE / "海外ETF期权及衍生策略产品研究0618_更新版.docx"

CN_FONT = "KaiTi"; EN_FONT = "Times New Roman"
INK = "000000"; GREY = "6B7280"; HEADER_FILL = "F2F4F7"; BORDER_COLOR = "D0D7DE"; YELLOW = "FFF2CC"

# ============================================================
# HELPERS
# ============================================================
def font_run(run, size=None, bold=None, color=None):
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = EN_FONT
    r_pr = run._element.get_or_add_rPr()
    rf = r_pr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        r_pr.append(rf)
    rf.set(qn("w:ascii"), EN_FONT); rf.set(qn("w:hAnsi"), EN_FONT); rf.set(qn("w:eastAsia"), CN_FONT)

def set_pf(p, size=11, before=0, after=6, line=1.2, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = line
    if align: p.alignment = align
    for r in p.runs: font_run(r, size=size)

def add_p(doc, ref, text, style="Normal", size=11, bold=False, color=INK, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text); font_run(run, size=size, bold=bold, color=color)
    set_pf(p, size=size, before=before, after=after, align=align)
    ref.addprevious(p._element); return p

def add_h(doc, ref, text, level):
    s = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    return add_p(doc, ref, text, s,
                 size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 before={1: 16, 2: 12, 3: 8}[level], after={1: 8, 2: 6, 3: 4}[level])

def add_src(doc, ref, text):
    return add_p(doc, ref, "资料来源：" + text, "资料来源", size=9, color=GREY, after=2)

def add_ttl(doc, ref, text):
    return add_p(doc, ref, text, "表格标题", size=10.5, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)

def fill_cell(cell, text, size=8.0, bold=False, shade=None):
    if shade: set_shade(cell, shade)
    cell.vertical_alignment = 1
    cell.text = ""
    p = cell.paragraphs[0]
    try: p.style = "表格正文"
    except: pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("" if text is None else str(text))
    font_run(r, size=size, bold=bold, color=INK)

def set_shade(cell, fill):
    tc = cell._tc.get_or_add_tcPr()
    s = tc.find(qn("w:shd"))
    if s is None: s = OxmlElement("w:shd"); tc.append(s)
    s.set(qn("w:fill"), fill)

def set_borders(table, color=BORDER_COLOR, size="4"):
    tbl = table._tbl.tblPr
    b = tbl.first_child_found_in("w:tblBorders")
    if b is None: b = OxmlElement("w:tblBorders"); tbl.append(b)
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = b.find(qn(f"w:{e}"))
        if el is None: el = OxmlElement(f"w:{e}"); b.append(el)
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color)

def add_tbl(doc, ref, headers, rows, fs=7.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, size=8.2, bold=True, shade=HEADER_FILL)
    for rd in rows:
        cells = table.add_row().cells
        for i, v in enumerate(rd):
            t = "" if v is None else str(v)
            fill_cell(cells[i], t, size=fs, shade=YELLOW if "待补充" in t else None)
    set_borders(table)
    ref.addprevious(table._element)
    return table

def remove_between(doc, start_text, end_text):
    body = doc.element.body
    children = list(body.iterchildren())
    si = ei = None
    for idx, child in enumerate(children):
        if isinstance(child, CT_P):
            t = "".join(child.itertext()).strip()
            if si is None and t.startswith(start_text): si = idx
            elif si is not None and t.startswith(end_text): ei = idx; break
    if si is None or ei is None:
        raise RuntimeError(f"Cannot find: '{start_text}' to '{end_text}'")
    ref = children[ei]
    for child in children[si:ei]: body.remove(child)
    return ref

# ============================================================
# CHAPTER 4 — CONTENT (all Chinese quotes use corner brackets only)
# ============================================================
def gen_ch4(doc, ref):
    add_h(doc, ref, "四、国内可行性分析", 1)

    # ---- opening ----
    add_p(doc, ref,
        "国内并非没有结构化产品的投资者基础。雪球类产品（Autocallable）在2023年末存续规模曾达约4000亿元，"
        "投资者对「敲入敲出」「票息」「保护垫」等概念并不陌生；在上交所，上证50ETF和沪深300ETF期权2025年全年"
        "日均成交量均超110万张，做市商体系成熟，近月平值合约流动性充裕。然而，国内至今未形成类似美国的"
        "期权收益型ETF和缓冲型ETF产品矩阵——市场规模和工具基础均已具备，但产品化始终未能落地。"
        "本章认为，其根本原因不在需求缺失，而在于政策导向、工具供给和产品化条件的结构性约束。"
        "国内公募期权策略产品应遵循「备兑增强优先推进、缓冲型中期储备、杠杆型可参考香港经验探索、"
        "高收益包装型和复杂结构化产品暂不适合」的优先序，分阶段推进。")

    # ---- 4.1 ----
    add_h(doc, ref, "4.1 政策环境：稳慎创新与风险管理优先", 2)

    add_p(doc, ref,
        "国内政策并非排斥衍生品创新，而是将衍生品定位为风险管理、套期保值、稳定投资行为和投资者保护的工具。"
        "海外市场更倾向于将期权收入、风险缓冲和目标结果产品作为可披露、可交易、由投资者自负盈亏的交易型工具；"
        "国内现阶段则强调衍生品服务于风险管理和中长期资金配置，产品宣传和结构设计均不宜将期权权利金包装为独立的高收益来源。")

    add_p(doc, ref,
        "从指数化投资政策看，2025年1月证监会印发的《促进资本市场指数化投资高质量发展行动方案》明确提出"
        "稳慎推进指数产品创新，在风险可测可控、投资者有效保护的前提下，研究推出多资产ETF、实物申赎模式"
        "跨市场债券ETF、银行间市场可转让指数基金等创新型指数产品，并研究拓展ETF底层资产类别。"
        "该方案同时提出持续丰富ETF期权、股指期货、股指期权等指数衍生品供给，为指数化投资提供更多风险管理工具。")

    add_p(doc, ref,
        "从公募基金政策看，2025年5月证监会印发的《推动公募基金高质量发展行动方案》提出制定公募基金参与"
        "金融衍生品投资指引，更好满足公募基金加强风险管理、稳定投资行为、丰富投资策略等需求，同时强调推动"
        "权益类基金产品创新、大力发展各类场内外指数基金，这意味着国内具备进一步研究期权增强、波动控制和"
        "目标风险类产品的政策空间。但从具体操作层面看，现行公募参与股票期权的约束仍包括权利金总额比例限制、"
        "未平仓合约面值限制、卖出认购期权须持有对应标的现货等，策略可用空间小于海外。")

    add_p(doc, ref,
        "从场外衍生品制度看，2026年5月15日证监会发布的《衍生品交易监督管理办法（试行）》进一步明确了"
        "证监会监管范围内场外衍生品的交易规则，并将于2026年11月16日起施行。该办法主要规范期货交易以外的"
        "互换合约、远期合约、非标准化期权合约及其组合交易，重点覆盖衍生品交易场所和以证券公司、期货公司"
        "为代表的衍生品经营机构。公募基金虽不是该办法的主要监管对象，但其作为场外衍生品的重要交易对手方，"
        "在对手方管理、风险监测和信息报送等方面将间接受益于更透明的制度环境。")

    add_p(doc, ref,
        "综合来看，国内政策正在为ETF期权及相关衍生策略产品的落地补齐工具和制度基础，但政策取向仍是稳慎创新和"
        "风险管理优先。因此，国内公募短期更适合研究低比例备兑增强（如备兑卖出比例控制在较低水平）、含权中低波动型产品"
        "和风险预算导向的资产配置产品；缓冲型ETF和目标结果ETF可作为中期储备方向；以期权权利金收入为主要收益来源的"
        "高收益增强型产品和高复杂度结构化收益产品暂不适合作为公募短期推进方向。")

    add_p(doc, ref,
        "值得注意的是，政策的「稳慎」取向并不等于「不做」。从2020年试点指引到2025年指数化投资方案再到2026年"
        "场外衍生品管理办法，政策脉络呈现清晰的渐进放开趋势——每一轮制度完善都在为下一阶段的产品创新预留空间。"
        "对于公募管理人而言，当前的关键不是在「能做」与「不能做」之间二选一，而是在既有制度框架内找到合规且有效的"
        "策略表达方式：控制期权覆盖比例而非全仓位卖出、强调波动控制而非高收益承诺、用指数化思维做策略而非用"
        "主动管理思维做择时。这种「在约束中找空间」的设计思路，是海外经验本土化的关键一步。")

    # policy table
    add_ttl(doc, ref, "表：国内衍生品相关政策文件梳理")
    add_tbl(doc, ref,
        ["时间", "政策文件或改革事项", "核心内容"],
        [
            ["2020年3月",
             "《证券期货经营机构参与股票期权交易试点指引》",
             "公募基金参与股票期权交易原则上以风险管理和套期保值为主要目的，并受到权利金总额、未平仓合约面值、备兑开仓标的持有等约束"],
            ["2025年1月",
             "《促进资本市场指数化投资高质量发展行动方案》",
             "稳慎推进指数产品创新，研究拓展ETF底层资产类别；持续丰富ETF期权、股指期货、股指期权等指数衍生品供给"],
            ["2025年5月",
             "《推动公募基金高质量发展行动方案》",
             "制定公募基金参与金融衍生品投资指引，推动权益类基金产品创新，大力发展各类场内外指数基金"],
            ["2026年5月发布\n(2026年11月施行)",
             "《衍生品交易监督管理办法（试行）》",
             "规范期货交易以外的互换、远期、非标准化期权及其组合交易，强化场外衍生品交易规则、风险监测和信息报送"],
        ], fs=7.8)
    add_src(doc, ref, "中国证监会、上交所、深交所、中金所、易方达产品研究。")

    # ---- 4.2 ----
    add_h(doc, ref, "4.2 为什么国内尚未形成规模化产品", 2)

    add_p(doc, ref,
        "国内尚未形成类似美国的option income ETF和Buffer ETF产品矩阵，并不是因为需求不存在，"
        "而是产品化条件仍不完整。对比海外经验与国内现状，核心约束体现在五个维度上。")

    add_p(doc, ref,
        "第一，监管定位存在本质差异。海外将option income ETF和Buffer ETF定位为可自由交易的交易所产品，"
        "由投资者自行判断和承担风险；国内将衍生品主要定位于风险管理和套期保值工具，"
        "监管对「将期权权利金包装为高收益来源」的叙事高度警惕。这一差异决定了国内产品的设计逻辑"
        "和营销口径必须与海外区分——不能将分派率等同于投资收益率，也不能将期权权利金包装为独立的收益引擎。")

    add_p(doc, ref,
        "第二，期权合约工具受限。海外Buffer ETF的核心工具是FLEX Options（灵活执行期权），"
        "允许管理人在交易所挂牌的标准化框架下，自定义到期日、行权价和期权类型，从而精确构建"
        "恰好一年的outcome period和精确的15%或30%下行缓冲边界。境内ETF期权合约仅有当月、下月"
        "及随后两个季月共四个月份，最长约9个月；行权价按交易所规则挂牌，通常5至10档。"
        "这使得海外经典的「一年期固定buffer+cap」结构在境内无法用场内工具精确复制。")

    add_p(doc, ref,
        "第三，公募基金参与存在量化约束。现行规则下，公募基金参与股票期权的权利金总额、"
        "未平仓合约面值以及卖出认购期权对应标的持有量等均存在比例限制。这些约束使得基金"
        "难以像JEPI那样将期权端的收入贡献提升至有意义的水平，期权的策略空间显著小于海外同类产品。")

    add_p(doc, ref,
        "第四，合约流动性高度集中。境内ETF期权的活跃合约集中在少数宽基ETF（上证50、沪深300、科创50等）"
        "的近月平值附近，远月合约和深度虚值合约流动性不足。对于需要持续滚动卖出期权、或在特定行权价"
        "构建保护边界的策略而言，流动性约束意味着更高的交易成本和滑点，削弱策略的净收益。")

    add_p(doc, ref,
        "第五，投资者适当性与销售适当性存在双重约束。期权交易本身有50万元资产门槛和知识测试要求；"
        "对于公募产品，监管要求产品风险收益特征清晰、适合目标投资者群体。将期权策略装入公募ETF后，"
        "如何确保零售投资者充分理解上行封顶、下行缓冲、波动拖累、路径依赖等非线性特征，"
        "是产品设计、信息披露和销售适当性管理需要系统性回答的问题。")

    add_p(doc, ref,
        "反观美国市场的成功条件，可以更清晰地看到国内差距的本质。美国期权策略ETF的爆发并非单一因素驱动，"
        "而是多个条件的共振：一是FLEX Options的交易所标准化框架，解决了定制化期权在公募产品中的估值、"
        "信息披露和流动性问题；二是SEC对「结果导向型」产品（Defined Outcome ETF）的注册和披露要求已形成"
        "成熟模板，管理人清楚合规边界；三是做市商生态完善——FLEX Options的报价和清算由OCC统一处理，"
        "做市商愿意为定制化合约提供双边报价；四是投资者教育和理财顾问渠道已将covered call和buffer策略"
        "纳入常规配置讨论，需求端有认知基础。这些条件国内目前仅部分具备——ETF期权流动性基础已打好，"
        "但定制化合約的制度框架、公募产品的估值和披露指引、以及理财渠道的策略认知，仍在渐进建设中。")

    # constraint table
    add_ttl(doc, ref, "表：国内期权策略产品化约束")
    add_tbl(doc, ref,
        ["约束维度", "国内现状", "对产品设计的影响"],
        [
            ["监管定位", "衍生品以风险管理和套期保值优先",
             "备兑策略应定位为低比例增强和风险控制工具，弱化高分红叙事"],
            ["合约工具", "标准化合约，期限9个月，行权价5-10档",
             "无法精确构建一年期outcome period和固定buffer/cap边界"],
            ["公募参与比例", "权利金总额、未平仓面值、备兑标的持有量均受限",
             "期权端的收入贡献难以提升至有意义的水平"],
            ["流动性分布", "活跃合约集中在近月平值，远月和虚值不足",
             "应围绕沪深300、上证50、科创50等流动性最好的底层做策略"],
            ["投资者适当性", "期权交易50万门槛，公募面向普通投资者",
             "产品设计须简单透明，非线性特征需充分披露"],
        ], fs=7.8)
    add_src(doc, ref, "境内交易所、监管文件、Cboe、易方达产品研究。")

    # ---- 4.3 ----
    add_h(doc, ref, "4.3 国内落地建议", 2)

    add_p(doc, ref,
        "基于海外产品格局和国内约束条件，国内公募期权策略产品的推进应按照"
        "「先简后繁、先大后小」的原则排序：优先推进海外已大规模验证、结构相对简单的备兑增强型产品；"
        "中期储备缓冲型产品，待场内或场外工具条件成熟后公募化；杠杆与反向型可参考香港经验探索推进；"
        "以高分红为核心卖点的option income产品和高复杂度结构化产品，与当前监管导向匹配度较低，"
        "短期内不适合作为公募推进方向。")

    # covered call
    add_p(doc, ref,
        "备兑增强型产品应作为短期最优先方向。从海外验证看，收益增强型是使用期权ETF中规模最大的类别，"
        "JEPI和JEPQ合计AUM近850亿美元，且策略逻辑透明——持有现货组合、卖出认购期权获取权利金。"
        "从国内条件看，上证50ETF和沪深300ETF期权2025年日均成交合计超220万张，做市商超过20家，"
        "近月平值合约买卖价差窄、市场深度好，足以支撑公募产品的日常备兑操作。"
        "招商证券关于红利ETF备兑策略的实证研究表明，备兑策略在A股震荡市中能够有效改善收益风险比——"
        "在窄幅震荡市中卖出虚值认购期权，权利金可增厚年化收益约2-4个百分点，而在趋势上涨市中仅小幅让渡上行空间。"
        "香港市场的实践也提供了邻近参照：国信证券对港美股备兑ETF产品线的全景梳理显示，"
        "香港已有数只以恒生指数和恒生科技指数为底层的备兑策略产品，产品规模和数量均在快速增长。")

    add_p(doc, ref,
        "产品设计上，底层标的应优先选择已有期权工具且流动性充裕的宽基ETF——沪深300ETF期权最为成熟，"
        "上证50ETF和科创50ETF次之，创业板ETF期权可作为差异化方向。行权策略建议卖出近月轻度虚值认购期权"
        "（Delta约0.2-0.3），覆盖比例建议初期控制在基金资产的10%-20%，以权利金增厚而非替代股息为核心目标。"
        "费率方面，考虑到备兑策略的主动管理属性高于纯被动ETF但低于主动权益基金，管理费率可设置在0.3%-0.5%。"
        "产品定位应明确为「权益收益增强」或「低波动权益配置」，目标投资者为对权益有配置需求、同时希望降低组合波动"
        "的中长期资金——包括养老金、理财子、保险资管和零售端的稳健型投资者。信息披露中须清晰展示权利金、"
        "股息和净值涨跌三者的关系，避免将月度分派等同于预期收益。")

    # buffer
    add_p(doc, ref,
        "缓冲型产品可作为中期储备方向。海外Buffer ETF约902亿美元规模，在财富管理渠道中验证了"
        "「用下行保护换取配置粘性」的产品逻辑——对于临近退休或对市场下跌敏感的投资者而言，"
        "牺牲部分上行空间换取明确的下行保护边界，比单纯的低波动策略更直观、更容易理解。"
        "招商证券的Buffer策略研究对FLEX Options构建保护结构的期权组合进行了详细拆解，"
        "并尝试在A股环境下模拟不同buffer深度（5%-20%）和期限组合的表现，初步验证了该策略在国内的适用潜力。"
        "但国内短期受限于合约期限（场内期权最长约9个月）和行权价（无法精确设定15%或30%边界），"
        "难以用场内工具精确复制经典的「一年期固定buffer+cap」结构。")

    add_p(doc, ref,
        "更现实的推进路径分为两步。第一步，在策略指数或基金专户层面，利用场外期权或收益互换"
        "模拟buffer结构，积累不同市场环境下的策略表现数据，完成策略在国内的「概念验证」——"
        "回答A股的buffer应该多深、期限应该多长、费用是否可接受等问题。第二步，"
        "待场内FLEX Options或类似定制化合约工具的制度条件成熟后，将已验证的策略平移至ETF框架。"
        "从政策节奏看，FLEX Options在国内的引入需要交易所、证监会和中国结算的系统性协调，"
        "预计至少需要2-3年以上的制度准备。因此缓冲型产品应定位为中期储备而非短期冲刺方向，"
        "公募管理人可在等待期内完成策略储备、投资者教育和合规预沟通。")

    # leveraged
    add_p(doc, ref,
        "杠杆与反向型产品可参考香港经验探索推进。香港市场已有29只杠杆/反向产品，"
        "SFC监管框架允许最高2倍杠杆且要求产品名称明确标注杠杆倍数。"
        "南方东英发行的CSOP SK Hynix Daily 2x Leveraged Product是香港规模最大的杠杆产品"
        "（AUM约108亿美元），验证了亚太市场对杠杆ETF的需求。内地投资者已可通过港股通间接参与，"
        "说明此类产品存在真实的交易和配置需求。此外，科创50ETF期权2026年以来日均成交量已进入前列，"
        "杠杆ETF的标的范围已从传统宽基大盘扩展到科技成长板块。"
        "但国内公募直接发行仍受净敞口限制和投资者适当性约束，短期内可优先开展产品研究和制度论证，而非直接推进发行。")

    # not priority
    add_p(doc, ref,
        "以高分红为核心卖点的option income产品、单股票杠杆ETF、0DTE策略产品和高复杂度结构化产品，"
        "与国内「稳慎创新」「投资者保护」的政策导向匹配度较低。其中，单股票杠杆ETF的个股集中风险和波动拖累效应"
        "对零售投资者尤为不利；将期权权利金包装为固定分派率的产品叙事在国内监管框架下难以获得认可。"
        "此类产品短期内不适合作为公募推进方向。Autocallable ETF（如CAIE）对国内雪球产品的公募化"
        "有长期参考价值——Calamos将Autocallable装入了每日可交易的ETF结构，证明了"
        "结构化产品与公募载体可以兼容——但需待场外衍生品监管框架和公募参与规则进一步成熟后再行评估。")

    # summary table
    add_ttl(doc, ref, "表：国内公募期权策略产品开发方向建议")
    add_tbl(doc, ref,
        ["开发方向", "优先级", "建议底层", "关键设计", "核心约束与前置条件"],
        [
            ["宽基备兑增强", "高",
             "沪深300、上证50、创业板",
             "低比例卖出近月轻度虚值认购期权；定位「权益增强」而非高分红",
             "期权覆盖比例、分派口径、投资者沟通口径须与监管对齐"],
            ["缓冲型产品", "中",
             "沪深300、中证1000",
             "利用场外期权或策略指数模拟buffer结构；先行回测验证",
             "合约期限和行权价定制化是公募化前提，需等FLEX Options或场外工具突破"],
            ["杠杆与反向型", "中",
             "宽基指数、港股通标的",
             "参考香港SFC监管框架，2倍杠杆，名称明确标注",
             "净敞口限制、投资者适当性、误用风险评估"],
            ["高收益option income", "低",
             "宽基或行业ETF",
             "不作为公募短期优先方向",
             "将期权权利金包装为高分红不符合国内监管导向"],
            ["单股票杠杆/反向", "低",
             "个股",
             "不作为公募短期优先方向",
             "个股集中风险、波动拖累、投资者误用风险均较高"],
        ], fs=7.6)
    add_src(doc, ref, "策略分析结论、监管材料、管理人产品资料、易方达产品研究。")

# ============================================================
# MAIN
# ============================================================
def main():
    print("=== 第四章重写（增强版）===")
    doc = Document(str(DOC_PATH))

    # Remove old Ch4 (between "四、" and "五、")
    ref = remove_between(doc, "四、", "五、")
    # Remove old Ch5 (between "五、" and "六、")
    ref = remove_between(doc, "五、", "六、")
    # Generate new Ch4 before old Ch6
    gen_ch4(doc, ref)

    doc.save(str(OUT_DOC))
    print(f"Saved: {OUT_DOC}")
    print("Done.")

if __name__ == "__main__":
    main()
