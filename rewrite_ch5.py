#!/usr/bin/env python3
"""
第五章回测设计重写 — 用实际回测结果替换空设计表
"""
import pandas as pd, numpy as np
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor
import warnings; warnings.filterwarnings('ignore')

BASE = Path("/Users/castle/Desktop/space for claude")
DOC_PATH = BASE / "海外ETF期权及衍生策略产品研究0618_更新版.docx"
OUT_DOC = BASE / "海外ETF期权及衍生策略产品研究0618_更新版.docx"

CN_FONT = "KaiTi"; EN_FONT = "Times New Roman"
INK = "000000"; GREY = "6B7280"; HEADER_FILL = "F2F4F7"; BORDER_COLOR = "D0D7DE"; YELLOW = "FFF2CC"

# ============================================================
# HELPERS (same as rewrite_ch4.py)
# ============================================================
def font_run(run, size=None, bold=None, color=None):
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = EN_FONT
    r_pr = run._element.get_or_add_rPr()
    rf = r_pr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        r_pr.append(rf)
    rf.set(qn("w:ascii"), EN_FONT); rf.set(qn("w:hAnsi"), EN_FONT); rf.set(qn("w:eastAsia"), CN_FONT)

def set_pf(p, size=11, before=0, after=6, line=1.2, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = line
    if align: p.alignment = align
    for r in p.runs: font_run(r, size=size)

def add_p(doc, ref, text, style="Normal", size=11, bold=False, color=INK, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text); font_run(run, size=size, bold=bold, color=color)
    set_pf(p, size=size, before=before, after=after, align=align)
    ref.addprevious(p._element); return p

def add_h(doc, ref, text, level):
    s = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    return add_p(doc, ref, text, s,
                 size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 before={1: 16, 2: 12, 3: 8}[level], after={1: 8, 2: 6, 3: 4}[level])

def add_src(doc, ref, text):
    return add_p(doc, ref, "资料来源：" + text, "资料来源", size=9, color=GREY, after=2)

def add_ttl(doc, ref, text):
    return add_p(doc, ref, text, "表格标题", size=10.5, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)

def fill_cell(cell, text, size=8.0, bold=False, shade=None):
    if shade: set_shade(cell, shade)
    cell.vertical_alignment = 1
    cell.text = ""
    p = cell.paragraphs[0]
    try: p.style = "表格正文"
    except: pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("" if text is None else str(text))
    font_run(r, size=size, bold=bold, color=INK)

def set_shade(cell, fill):
    tc = cell._tc.get_or_add_tcPr()
    s = tc.find(qn("w:shd"))
    if s is None: s = OxmlElement("w:shd"); tc.append(s)
    s.set(qn("w:fill"), fill)

def set_borders(table, color=BORDER_COLOR, size="4"):
    tbl = table._tbl.tblPr
    b = tbl.first_child_found_in("w:tblBorders")
    if b is None: b = OxmlElement("w:tblBorders"); tbl.append(b)
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = b.find(qn(f"w:{e}"))
        if el is None: el = OxmlElement(f"w:{e}"); b.append(el)
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color)

def add_tbl(doc, ref, headers, rows, fs=7.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, size=8.2, bold=True, shade=HEADER_FILL)
    for rd in rows:
        cells = table.add_row().cells
        for i, v in enumerate(rd):
            t = "" if v is None else str(v)
            fill_cell(cells[i], t, size=fs, shade=YELLOW if "待补充" in t else None)
    set_borders(table)
    ref.addprevious(table._element)
    return table

def remove_between(doc, start_text, end_text):
    body = doc.element.body
    children = list(body.iterchildren())
    si = ei = None
    for idx, child in enumerate(children):
        if isinstance(child, CT_P):
            t = "".join(child.itertext()).strip()
            if si is None and t.startswith(start_text): si = idx
            elif si is not None and t.startswith(end_text): ei = idx; break
    if si is None or ei is None:
        raise RuntimeError(f"Cannot find: '{start_text}' to '{end_text}'")
    ref = children[ei]
    for child in children[si:ei]: body.remove(child)
    return ref

# ============================================================
# CHAPTER 5
# ============================================================
def gen_ch5(doc, ref):
    add_h(doc, ref, "五、回测设计", 1)

    add_p(doc, ref,
        "回测要回答的核心问题是：海外成熟策略落到国内市场后，是否真的能改善投资者体验。"
        "评价标准不应只看收益率，还须同时观察波动率、最大回撤、月度胜率和风险调整后收益。"
        "以下以50ETF为底层、以上交所ETF期权上市以来的历史数据为基础，"
        "对月度备兑认购策略进行回测。期权定价采用Black-Scholes模型，"
        "隐含波动率取自上交所50ETF期权VIX指数（iVIX），无风险利率取2.5%。")

    # ---- 5.1 ----
    add_h(doc, ref, "5.1 备兑认购策略回测", 2)

    add_p(doc, ref,
        "策略设定：每月第三个周五（或最近交易日）卖出当月到期、行权价等于或略高于ETF现价的"
        "认购期权，持有至到期；到期日平掉到期仓位并开立下一个月的仓位。标的为上证50ETF（510050），"
        "回测区间自2015年2月50ETF期权上市至2026年5月，共137个月度周期。"
        "分别测试平值（ATM）、虚值2%（OTM 102%）和虚值3%（OTM 103%）三种行权价选择。")

    # Results table
    add_ttl(doc, ref, "表：50ETF月度备兑认购策略回测结果（2015.2—2026.5）")
    add_tbl(doc, ref,
        ["策略", "年化回报", "年化波动", "最大回撤", "Sharpe", "Calmar", "Sortino", "月度胜率"],
        [
            ["买入持有(同期)", "5.6%", "18.8%", "-40.8%", "0.16", "0.14", "0.25", "51.1%"],
            ["备兑ATM", "9.5%", "10.1%", "-20.9%", "0.69", "0.45", "0.62", "76.6%"],
            ["备兑OTM(102%)", "9.7%", "12.2%", "-23.7%", "0.59", "0.41", "0.64", "67.9%"],
            ["备兑OTM(103%)", "9.4%", "13.2%", "-24.7%", "0.53", "0.38", "0.62", "66.4%"],
        ], fs=7.8)

    add_p(doc, ref,
        "回测结果显示，月度备兑策略在所有风险调整后指标上均显著优于买入持有。"
        "ATM备兑的年化回报从5.6%提升至9.5%，同时年化波动率从18.8%降至10.1%、"
        "最大回撤从-40.8%降至-20.9%，Sharpe Ratio从0.16提升至0.69。"
        "月均权利金收入约2.5%（ATM），年化权利金收入约30%，但其中相当部分被标的下跌和"
        "上行封顶所抵消，实际净增厚约3-4个百分点。为审慎起见，考虑实际交易中买卖价差、"
        "滑点和佣金的影响（估计合计约0.3%-0.5%每月），实际净增厚可能在2-3个百分点。")

    add_p(doc, ref,
        "从行权价选择看，ATM备兑在Sharpe Ratio（0.69）和波动控制（10.1%）上表现最佳；"
        "OTM 102%的年化回报略高（9.7%）但波动也更高（12.2%）。"
        "这是因为ATM合约权利金更高（月均2.5% vs OTM的1.6%），对下行缓冲更充分，"
        "而OTM合约保留了更多上行空间但权利金收入较少。对于追求风险调整后收益的产品设计，"
        "ATM或轻度虚值（Delta约0.5-0.6）是更合适的选择。")

    # Yearly breakdown
    add_ttl(doc, ref, "表：ATM备兑策略分年表现")
    add_tbl(doc, ref,
        ["年份", "ETF年化回报", "备兑年化回报", "月均权利金", "备兑月度胜率", "ETF月度胜率"],
        [
            ["2015", "-5.7%", "-3.3%", "4.1%", "66.7%", "50.0%"],
            ["2016", "15.5%", "18.9%", "2.8%", "75.0%", "66.7%"],
            ["2017", "29.0%", "16.3%", "1.6%", "100.0%", "66.7%"],
            ["2018", "-22.2%", "5.1%", "2.7%", "75.0%", "25.0%"],
            ["2019", "28.0%", "19.9%", "2.5%", "75.0%", "75.0%"],
            ["2020", "26.1%", "20.1%", "2.8%", "83.3%", "75.0%"],
            ["2021", "-16.8%", "-3.0%", "2.5%", "66.7%", "33.3%"],
            ["2022", "-5.8%", "-1.4%", "2.5%", "75.0%", "41.7%"],
            ["2023", "-20.4%", "-0.4%", "2.1%", "58.3%", "16.7%"],
            ["2024", "19.4%", "12.6%", "2.1%", "75.0%", "50.0%"],
            ["2025", "21.6%", "18.9%", "2.1%", "91.7%", "75.0%"],
            ["2026(至5月)", "-11.9%", "10.3%", "2.0%", "80.0%", "20.0%"],
        ], fs=7.6)

    add_p(doc, ref,
        "分年数据揭示了备兑策略最核心的价值——在下行市中提供保护缓冲。"
        "2018年（贸易摩擦）ETF全年下跌22.2%，备兑策略逆势录得+5.1%；"
        "2023年（弱复苏+外资流出）ETF下跌20.4%，备兑仅微亏0.4%；"
        "2021年（结构性熊市）ETF跌16.8%，备兑仅跌3.0%。"
        "代价是在强趋势上涨市中让渡部分上行空间——2017年ETF涨29.0%，备兑涨16.3%，"
        "少赚约13个百分点。这种「以有限上行空间换取实质性下行保护」的收益特征，"
        "正是备兑策略在震荡市和温和下跌市中改善投资者体验的核心逻辑。")

    add_src(doc, ref,
        "上证50ETF历史行情（Yahoo Finance，510050.SS），50ETF期权VIX指数（akshare/上交所），"
        "Black-Scholes期权定价，无风险利率取2.5%。回测未计入买卖价差和佣金，"
        "实际策略表现可能略低于理论值。易方达产品研究。")

    # ---- 5.2 Buffer ----
    add_h(doc, ref, "5.2 Buffer策略回测", 2)

    add_p(doc, ref,
        "Buffer策略的回测与备兑策略有本质不同：它要求精确构建特定到期日和行权价的期权组合"
        "（买入保护性认沽+卖出更低行权价认沽+卖出上方认购期权），"
        "而境内标准化期权无法提供FLEX Options的到期日和行权价灵活性。"
        "因此，Buffer策略回测的目的并非追求精确的历史模拟，而是回答两个更基础的问题："
        "第一，在境内现有合约约束下，能否拼出近似buffer的效果？"
        "第二，buffer结构的保护成本是否在产品化可接受的范围内？")

    add_p(doc, ref,
        "回测设定：以沪深300指数为底层，采用Black-Scholes模型对期权组合进行定价——"
        "买入平值保护性认沽、卖出buffer边界认沽（行权价=标的×（1-buffer%））、"
        "卖出5%虚值认购期权以补贴保护成本。隐含波动率取自上交所300ETF期权VIX指数。"
        "回测区间从2021年3月（yfinance可获取的沪深300指数数据起点）至2026年5月，"
        "每季度换仓一次，buffer深度测试2%至10%共五个档位。")

    add_ttl(doc, ref, "表：Buffer策略回测设定")
    add_tbl(doc, ref,
        ["项目", "设定"],
        [
            ["底层指数", "沪深300（000300.SH）"],
            ["回测区间", "2021年3月11日—2026年5月31日，21个季度周期"],
            ["期权定价", "Black-Scholes模型，IV取300ETF期权VIX指数"],
            ["换仓频率", "每季度一次（季月合约）"],
            ["Buffer深度", "2%、4%、6%、8%、10%"],
            ["构建方式", "买入ATM保护性认沽 + 卖出buffer边界认沽 + 卖出5%虚值认购补贴成本"],
            ["评价指标", "年化回报、年化波动、最大回撤、Sharpe、对冲成本率、突破buffer季度数"],
        ], fs=7.8)

    add_ttl(doc, ref, "表：不同Buffer深度的回测结果（2021.3—2026.5）")
    add_tbl(doc, ref,
        ["Buffer深度", "年化回报", "年化波动", "最大回撤", "Sharpe", "年化对冲成本", "突破Buffer季度数"],
        [
            ["无Buffer(买入持有)", "0.5%", "32.8%", "-38.6%", "-0.06", "—", "—"],
            ["2%", "4.6%", "20.5%", "-19.5%", "0.10", "-1.2%（权利金覆盖）", "7"],
            ["4%", "3.8%", "17.9%", "-17.2%", "0.07", "-0.4%（权利金覆盖）", "6"],
            ["6%", "3.8%", "15.7%", "-14.7%", "0.08", "0.1%", "3"],
            ["8%", "2.4%", "13.7%", "-13.1%", "0.00", "0.5%", "3"],
            ["10%", "2.4%", "11.9%", "-11.4%", "-0.01", "0.8%", "2"],
        ], fs=7.6)

    add_p(doc, ref,
        "回测结果呈现两个值得关注的特征。第一，buffer结构在降低波动和回撤方面效果明确——"
        "10% buffer将年化波动从32.8%降至11.9%、最大回撤从-38.6%降至-11.4%，"
        "突破buffer边界的季度从7个降至2个。第二，与直觉相反，浅buffer（2%-6%）的"
        "对冲成本极低甚至为负——卖出的虚值认购期权权利金足以覆盖保护性认沽的净成本，"
        "策略回报反而高于买入持有。其原因是回测期内沪深300指数整体震荡、隐含波动率"
        "处于中高水平（13%-46%），认购期权权利金较为丰厚。")

    add_p(doc, ref,
        "但这一结果的适用性需要审慎评估。首先，回测仅覆盖2021年3月至2026年5月共21个季度，"
        "期间沪深300指数整体持平（年化仅0.5%），buffer策略在震荡市中的表现天然优于"
        "买入持有，但在强趋势牛市中可能显著落后。其次，上述回测假设可以按任意行权价构建"
        "期权组合，而境内沪深300股指期权的行权价按交易所规则挂牌，实际可用的行权价"
        "与理论buffer边界之间存在偏差，将对冲成本和保护效果产生额外影响。"
        "这也印证了第四章的判断：Buffer ETF在国内的落地需要等待FLEX Options"
        "或类似定制化合约工具的制度突破。")

    add_src(doc, ref,
        "沪深300指数行情（Yahoo Finance，000300.SS），300ETF期权VIX指数（akshare/上交所），"
        "Black-Scholes期权定价，无风险利率取2.5%。回测未计入买卖价差和滑点，"
        "且假设可按任意行权价构建期权组合（实际情况受交易所挂牌行权价限制），"
        "实际对冲成本可能更高、保护效果可能低于理论值。易方达产品研究。")

# ============================================================
# MAIN
# ============================================================
def main():
    print("=== 第五章回测重写 ===")
    doc = Document(str(DOC_PATH))

    # Find "五、回测设计" and remove old Ch5 content
    body = doc.element.body
    children = list(body.iterchildren())
    si = None
    sectPr = None
    for idx, child in enumerate(children):
        # Save sectPr before removing
        if child.tag == qn("w:sectPr"):
            sectPr = child
        if isinstance(child, CT_P):
            t = "".join(child.itertext()).strip()
            if t.startswith("六、回测设计") or t.startswith("五、回测设计"):
                si = idx
                break
    if si is None:
        raise RuntimeError("Cannot find 五、回测设计")

    # Remove old Ch5 content (but NOT sectPr)
    for child in children[si:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    # Now generate new Ch5 — need a ref before sectPr
    # Insert a sentinel paragraph before sectPr
    sentinel = doc.add_paragraph()
    if sectPr is not None:
        sectPr.addprevious(sentinel._element)
    gen_ch5(doc, sentinel._element)
    body.remove(sentinel._element)

    doc.save(str(OUT_DOC))
    print("Saved — Chapter 5 backtest rewritten with real data.")

if __name__ == "__main__":
    main()
