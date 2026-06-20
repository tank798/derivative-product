#!/usr/bin/env python3
"""
海外ETF期权及衍生策略产品研究 - 第二章、第三章重写脚本 V2
11项修正：AUM%修复、产品发行表格化、资金流拆解、删除效率/产品矩阵/2.6、
          费率分析迁入Ch3、中文引号「」、小标题仅管理人名称
"""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
import warnings
warnings.filterwarnings('ignore')

# ============================================================
BASE = Path("/Users/castle/Desktop/space for claude")
RAW_PATH = BASE / "global derivative ETFs-20260616.xlsx"
WB_PATH = BASE / "海外ETF期权及衍生策略产品研究_核验底稿.xlsx"
DOC_PATH = BASE / "海外ETF期权及衍生策略产品研究0618.docx"
OUT_DOC = BASE / "海外ETF期权及衍生策略产品研究0618_更新版.docx"
OUT_WB = BASE / "海外ETF期权及衍生策略产品研究_核验底稿_更新版.xlsx"

CN_FONT = "KaiTi"
EN_FONT = "Times New Roman"
INK = "000000"
GREY = "6B7280"
HEADER_FILL = "F2F4F7"
BORDER_COLOR = "D0D7DE"
YELLOW = "FFF2CC"

TYPE_ORDER = ["收益增强型", "风险缓冲型", "杠杆与反向型", "另类衍生策略型"]
TYPE_SHORT = {"收益增强型": "收益增强", "风险缓冲型": "风险缓冲",
              "杠杆与反向型": "杠杆反向", "另类衍生策略型": "另类策略"}

# Verified benchmark ETF fees (percentage points, same unit as Morningstar)
BENCHMARKS = {
    'S&P 500':      {'ticker': 'IVV', 'fee': 0.03, 'name': 'iShares Core S&P 500 ETF'},
    'Nasdaq 100':   {'ticker': 'QQQM','fee': 0.15, 'name': 'Invesco Nasdaq 100 ETF'},
    'Russell 2000': {'ticker': 'IWM', 'fee': 0.19, 'name': 'iShares Russell 2000 ETF'},
    'Dow Jones':    {'ticker': 'DIA', 'fee': 0.16, 'name': 'SPDR Dow Jones ETF'},
}

# ============================================================
# DATA
# ============================================================
def load_data():
    raw = pd.read_excel(RAW_PATH, sheet_name="20260616_global_derivative")
    raw['SecId'] = raw['SecId'].astype(str).str.strip()
    wb = pd.read_excel(WB_PATH, sheet_name="衍生策略分类明细")
    wb['SecId'] = wb['SecId'].astype(str).str.strip()

    raw['使用期权'] = raw['SecId'].map(wb.set_index('SecId')['是否使用期权']).fillna('否')
    raw['产品类型'] = raw['SecId'].map(wb.set_index('SecId')['产品类型']).fillna('另类衍生策略型')

    raw['Fund Size USD'] = pd.to_numeric(raw['Fund Size USD'], errors='coerce').fillna(0)
    raw['Fee'] = pd.to_numeric(raw['Annual Report Net Expense Ratio'], errors='coerce')
    raw['Inception Date'] = pd.to_datetime(raw['Inception Date'], errors='coerce')
    raw['Inception Year'] = raw['Inception Date'].dt.year.fillna(0).astype(int)

    flow_map = {
        'Flow_1Mo': 'Est Fund-Level Net Flow 1 Mo (Mo-End) USD',
        'Flow_3Mo': 'Est Fund-Level Net Flow 3 Mo (Mo-End) USD',
        'Flow_6Mo': 'Est Fund-Level Net Flow 6 Mo (Mo-End) USD',
        'Flow_YTD': 'Est Fund-Level Net Flow YTD (Mo-End) USD',
        'Flow_1Yr': 'Est Fund-Level Net Flow 1 Yr (Mo-End) USD',
        'Flow_3Yr': 'Est Fund-Level Net Flow 3 Yr (Mo-End) USD',
    }
    for k, v in flow_map.items():
        raw[k] = pd.to_numeric(raw[v], errors='coerce')

    perf_map = {
        'Ret_1Yr': 'Ret 1 Yr (Mo-End)', 'Ret_3Yr': 'Ret Annlzd 3 Yr (Mo-End)',
        'StdDev_1Yr': 'Std Dev 1 Yr (Mo-End) Risk Currency',
        'StdDev_3Yr': 'Std Dev 3 Yr (Mo-End) Risk Currency',
        'Sharpe_1Yr': 'Sharpe Ratio 1 Yr (Mo-End) Risk Currency',
        'Sharpe_3Yr': 'Sharpe Ratio 3 Yr (Mo-End) Risk Currency',
        'Beta_1Yr': 'Beta 1 Yr (Mo-End) Risk Currency',
        'Beta_3Yr': 'Beta 3 Yr (Mo-End) Risk Currency',
    }
    for k, v in perf_map.items():
        raw[k] = pd.to_numeric(raw[v], errors='coerce')

    return raw

def get_opt(raw):
    return raw[raw['使用期权'] == '是'].copy()

def get_perf(raw):
    """Representative product performance lookup"""
    targets = ['JEPI','JEPQ','QYLD','XYLD','SPYI','QQQI','BUFR','BALT','BUFB','BUFF',
               'TQQQ','SQQQ','SVOL','DBMF','NTSX','TAIL','CTA','VIXY','RYLD','DIVO']
    p = raw[raw['Ticker'].isin(targets)].copy()
    p['Fund Size USD'] = pd.to_numeric(p['Fund Size USD'], errors='coerce').fillna(0)
    p = p.sort_values('Fund Size USD', ascending=False).drop_duplicates(subset=['Ticker'], keep='first')
    return p.set_index('Ticker')

def classify_benchmark(row):
    t = str(row['Ticker']).upper()
    n = str(row['Name']).lower()
    if t in ['JEPI','XYLD','SPYI','HELO','BALT','BUFR','BUFP','KNG','FTHI','GPIX']: return 'S&P 500'
    if t in ['JEPQ','QYLD','QQQI']: return 'Nasdaq 100'
    if t in ['RYLD','IWMI']: return 'Russell 2000'
    if t in ['DIVO','DJIA']: return 'Dow Jones'
    if 's&p 500' in n: return 'S&P 500'
    if 'nasdaq' in n and '100' in n: return 'Nasdaq 100'
    if 'russell 2000' in n: return 'Russell 2000'
    if 'dow' in n: return 'Dow Jones'
    return None

def fmt_f(val):
    if pd.isna(val): return "N/A"
    return f"{val:.2f}%"

def fmt_r(val):
    if pd.isna(val): return "N/A"
    return f"{val:.2f}%"

def fmt_a(val):
    if pd.isna(val) or val == 0: return "N/A"
    return f"{val:.2f}"

def fmt_bp(val):
    """Format as basis points"""
    return f"{val*100:.0f}bp"

# ============================================================
# WORD HELPERS
# ============================================================
def font_run(run, size=None, bold=None, color=None):
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color is not None: run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = EN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts if r_pr.rFonts is not None else OxmlElement("w:rFonts")
    if r_pr.rFonts is None: r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), EN_FONT)
    r_fonts.set(qn("w:hAnsi"), EN_FONT)
    r_fonts.set(qn("w:eastAsia"), CN_FONT)

def set_pf(paragraph, size=11, before=0, after=6, line=1.2, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    if align is not None: paragraph.alignment = align
    for run in paragraph.runs:
        font_run(run, size=size)

def add_p(doc, ref, text, style_name="Normal", size=11, bold=False, color=INK, before=0, after=6, align=None):
    paragraph = doc.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    font_run(run, size=size, bold=bold, color=color)
    set_pf(paragraph, size=size, before=before, after=after, align=align)
    ref.addprevious(paragraph._element)
    return paragraph

def add_h(doc, ref, text, level):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    return add_p(doc, ref, text, style,
                 size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 before={1: 16, 2: 12, 3: 8}[level], after={1: 8, 2: 6, 3: 4}[level])

def add_src(doc, ref, text):
    return add_p(doc, ref, "资料来源：" + text, "资料来源", size=9, color=GREY, after=2)

def add_note(doc, ref, text):
    return add_p(doc, ref, "注：" + text, "注释", size=9, color=GREY, after=6)

def add_ttl(doc, ref, text):
    return add_p(doc, ref, text, "表格标题", size=10.5, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)

def fill_cell(cell, text, size=8.0, bold=False, shade=None):
    if shade: set_cell_shading(cell, shade)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    try: p.style = "表格正文"
    except KeyError: pass
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("" if text is None else str(text))
    font_run(run, size=size, bold=bold, color=INK)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_borders(table, color=BORDER_COLOR, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None: borders = OxmlElement("w:tblBorders"); tbl_pr.append(borders)
    for edge in ("top","left","bottom","right","insideH","insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None: el = OxmlElement(tag); borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)

def add_tbl(doc, ref, headers, rows, font_size=7.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        fill_cell(table.rows[0].cells[idx], header, size=8.2, bold=True, shade=HEADER_FILL)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            text = "" if value is None else str(value)
            shade = YELLOW if "待补充" in str(text) else None
            fill_cell(cells[idx], text, size=font_size, shade=shade)
    set_borders(table)
    ref.addprevious(table._element)
    return table

def remove_between_headings(doc, start_text, end_text):
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = end_idx = None
    for idx, child in enumerate(children):
        if isinstance(child, CT_P):
            text = "".join(child.itertext()).strip()
            if start_idx is None and text.startswith(start_text): start_idx = idx
            elif start_idx is not None and text.startswith(end_text): end_idx = idx; break
    if start_idx is None or end_idx is None:
        raise RuntimeError(f"Cannot find: {start_text} to {end_text}")
    ref = children[end_idx]
    for child in children[start_idx:end_idx]: body.remove(child)
    return ref

# ============================================================
# MANAGER RANKING
# ============================================================
def build_manager_ranking(opt):
    mgr = opt.groupby('Firm Name').agg(
        产品数量=('SecId', 'count'),
        总AUM亿=('Fund Size USD', lambda x: x.sum()/1e8),
    ).reset_index().sort_values('总AUM亿', ascending=False)

    for i, row in mgr.iterrows():
        sub = opt[opt['Firm Name'] == row['Firm Name']]
        sub_fee = sub[sub['Fee'].notna()]
        if len(sub_fee) > 0 and sub_fee['Fund Size USD'].sum() > 0:
            mgr.at[i, 'AUM加权费率'] = (sub_fee['Fee'] * sub_fee['Fund Size USD']).sum() / sub_fee['Fund Size USD'].sum()
        for typ in TYPE_ORDER:
            mgr.at[i, typ+'_数量'] = int((sub['产品类型'] == typ).sum())
    return mgr

# ============================================================
# CHAPTER 2
# ============================================================
def generate_chapter2(doc, ref, opt, mgr, perf_lu):
    top5 = mgr.head(5)
    top5_firms = top5['Firm Name'].tolist()
    total_opt_aum_亿 = opt['Fund Size USD'].sum() / 1e8  # FIXED: correct denominator

    add_h(doc, ref, "二、管理人布局与竞争优势", 1)

    # ---- Opening ----
    add_p(doc, ref,
        "在使用期权的衍生策略ETF领域，管理人竞争格局呈现高度集中与深度差异化并存的特征。"
        "按AUM排序，前五大管理人分别为JPMorgan、First Trust、Innovator ETFs、NEOS和Global X，"
        f"五家合计AUM约{top5['总AUM亿'].sum():.0f}亿美元，占全球使用期权ETF总规模的"
        f"{top5['总AUM亿'].sum()/total_opt_aum_亿*100:.1f}%。")

    # Manager overview table (FIXED AUM%)
    add_ttl(doc, ref, "表：使用期权的衍生策略ETF头部管理人总览")
    overview_rows = []
    for idx, (_, r) in enumerate(top5.iterrows()):
        pct = r['总AUM亿'] / total_opt_aum_亿 * 100  # FIXED
        aum_w_fee = fmt_f(r.get('AUM加权费率'))
        types_parts = []
        for typ in TYPE_ORDER:
            cnt = int(r.get(f'{typ}_数量', 0))
            if cnt > 0: types_parts.append(f"{TYPE_SHORT[typ]}{cnt}只")
        overview_rows.append([
            idx+1, r['Firm Name'], int(r['产品数量']),
            f"{r['总AUM亿']:.2f}", f"{pct:.1f}%",
            aum_w_fee, " ".join(types_parts),
        ])
    add_tbl(doc, ref,
        ["排名", "管理人", "产品数量", "AUM（亿美元）", "占使用期权ETF AUM",
         "AUM加权费率", "策略覆盖"],
        overview_rows, font_size=7.8)
    add_src(doc, ref,
        "Morningstar、管理人产品资料、易方达产品研究。数据截至2026年5月31日，使用期权产品按二级复核后分类。")
    add_note(doc, ref,
        "AUM加权费率 = Σ（各产品费率×AUM）/ Σ（AUM），反映投资者实际承担的费率水平。"
        f"全球使用期权ETF总AUM约{total_opt_aum_亿:.2f}亿美元。")

    # Strategic positioning (moved from old 2.6)
    jpm_n = len(opt[opt['Firm Name'] == 'JPMorgan'])
    neos_n = len(opt[opt['Firm Name'] == 'Neos Funds'])
    ft_n = len(opt[opt['Firm Name'] == 'First Trust'])
    innov_n = len(opt[opt['Firm Name'] == 'Innovator ETFs'])
    gx_n = len(opt[opt['Firm Name'] == 'Global X Funds'])
    gx_aum = opt[opt['Firm Name'] == 'Global X Funds']['Fund Size USD'].sum() / 1e8
    add_p(doc, ref,
        "从策略路径看，五家管理人呈现出三种差异化的定位："
        f"JPMorgan和NEOS以少数核心产品驱动规模——JPMorgan仅{jpm_n}只产品达到896亿美元，"
        f"Top3产品AUM集中度达98.9%；NEOS仅{neos_n}只产品达到249亿美元，Top3集中度96.9%。"
        f"First Trust和Innovator则以产品矩阵的广度取胜——分别以{ft_n}只和{innov_n}只产品覆盖不同月份、"
        "缓冲深度和底层资产，Top3集中度仅为27.8%和15.9%。"
        "Global X作为最早布局covered call ETF的管理人（QYLD/XYLD均于2013年发行），"
        "虽拥有显著的先发时间窗口，但后续产品线扩张速度明显落后，"
        f"{gx_n}只产品仅贡献{gx_aum:.0f}亿美元AUM，且近1年出现资金净流出。")

    # ---- Per-manager sections ----
    for i, firm in enumerate(top5_firms):
        sub = opt[opt['Firm Name'] == firm]
        firm_aum = sub['Fund Size USD'].sum() / 1e8
        firm_n = len(sub)

        add_h(doc, ref, f"2.{i+1} {firm}", 2)

        # Sub-fee for fee-weighted calc
        sub_fee = sub[sub['Fee'].notna()]
        wfee = (sub_fee['Fee'] * sub_fee['Fund Size USD']).sum() / sub_fee['Fund Size USD'].sum() if len(sub_fee) > 0 and sub_fee['Fund Size USD'].sum() > 0 else None

        # Brief intro
        type_counts = sub['产品类型'].value_counts()
        type_desc = "、".join([f"{TYPE_SHORT[t]}{type_counts.get(t,0)}只" for t in TYPE_ORDER if type_counts.get(t,0) > 0])
        fee_str = f"，AUM加权费率为{fmt_f(wfee)}" if wfee else ""

        intro = f"{firm}在使用期权的衍生策略ETF中AUM约{firm_aum:.2f}亿美元，产品数量{firm_n}只，覆盖{type_desc}{fee_str}。"
        add_p(doc, ref, intro)

        # ---- 1) Product launch table ----
        add_ttl(doc, ref, f"表：{firm}产品发行节奏——按产品类型×年份")

        # Build cross table
        sub_cross = sub[sub['Inception Year'].between(2005, 2026)]
        years = sorted(sub_cross['Inception Year'].unique())
        # Determine which types are present
        present_types = [t for t in TYPE_ORDER if type_counts.get(t, 0) > 0]

        launch_rows = []
        for yr in years:
            yr_data = sub_cross[sub_cross['Inception Year'] == yr]
            row_data = [str(yr)]
            for typ in present_types:
                cnt = int((yr_data['产品类型'] == typ).sum())
                row_data.append(str(cnt) if cnt > 0 else "-")
            row_data.append(str(len(yr_data)))
            launch_rows.append(row_data)

        add_tbl(doc, ref,
            ["年份"] + [TYPE_SHORT[t] for t in present_types] + ["合计"],
            launch_rows, font_size=8.0)
        add_src(doc, ref, "Morningstar、管理人产品资料、易方达产品研究。仅统计当前存续的使用期权产品。")

        # Summary paragraph
        # Group years into phases
        all_years = sorted(sub_cross['Inception Year'].unique())
        summary = build_launch_summary(firm, sub_cross, present_types, all_years)
        add_p(doc, ref, summary)

        # ---- 2) Net flows breakdown ----
        add_ttl(doc, ref, f"表：{firm}近期资金流向")

        flow_labels = [
            ('Flow_1Mo', '近1个月'), ('Flow_3Mo', '近3个月'),
            ('Flow_6Mo', '近6个月'), ('Flow_YTD', '年初至今'),
            ('Flow_1Yr', '近1年'), ('Flow_3Yr', '近3年'),
        ]

        flow_rows = []
        for fl_col, fl_label in flow_labels:
            has_data = sub[sub[fl_col].notna()]
            n_data = len(has_data)
            n_missing = len(sub) - n_data
            total_flow = has_data[fl_col].sum() / 1e8
            n_in = int((has_data[fl_col] > 0).sum())
            n_out = int((has_data[fl_col] < 0).sum())
            inflow_sum = has_data[has_data[fl_col] > 0][fl_col].sum() / 1e8
            outflow_sum = has_data[has_data[fl_col] < 0][fl_col].sum() / 1e8

            # Top contributor
            if n_data > 0:
                top_idx = has_data[fl_col].abs().nlargest(1).index[0]
                top_row = has_data.loc[top_idx]
                top_str = f"{top_row['Ticker']}（{top_row[fl_col]/1e8:+.2f}亿）"
            else:
                top_str = "无数据"

            missing_note = f"（缺{n_missing}只）" if n_missing > 0 else ""
            flow_rows.append([
                fl_label, f"{total_flow:+.2f}",
                f"流入{n_in}只/流出{n_out}只{missing_note}",
                f"流入合计+{inflow_sum:.2f}亿 / 流出合计{outflow_sum:.2f}亿",
                top_str,
            ])

        add_tbl(doc, ref,
            ["时间窗口", "净流量合计（亿美元）", "产品分布", "流入/流出明细", "最大贡献产品"],
            flow_rows, font_size=7.8)
        add_src(doc, ref, "Morningstar资金流数据。净流量为管理人旗下所有使用期权产品对应时间窗口的资金净流入之和。")
        # Count products with Flow_3Yr data per manager
        fl_counts = {}
        for mgr in ['JPMorgan', 'First Trust', 'Innovator ETFs', 'Neos Funds', 'Global X Funds']:
            sub = opt[opt['Firm Name'] == mgr]
            with_flow = sub[sub['Flow_3Yr'].notna()]
            fl_counts[mgr] = f"{len(with_flow)}/{len(sub)}"
        add_note(doc, ref,
            f"Flow_3Yr仅覆盖成立满3年的产品（JPMorgan {fl_counts['JPMorgan']}只、"
            f"First Trust {fl_counts['First Trust']}只、Innovator {fl_counts['Innovator ETFs']}只、"
            f"NEOS {fl_counts['Neos Funds']}只、Global X {fl_counts['Global X Funds']}只），"
            "与较短时间窗口数据不可直接横向对比。"
            "流入/流出明细为各产品正/负流向各自加总，两者相抵后等于净流量合计。")


def build_launch_summary(firm, sub, present_types, all_years):
    """Generate a natural-language summary of the product launch history"""
    if len(all_years) == 0:
        return ""

    min_yr = min(all_years)
    max_yr = max(all_years)

    # Count total by year
    yr_counts = sub.groupby('Inception Year').size()

    # Determine the dominant type per year group
    parts = []

    # Early phase: first few years
    early_yrs = [y for y in all_years if y <= min_yr + 2] if len(all_years) > 3 else all_years[:max(1, len(all_years)//3)]
    if early_yrs:
        early_sub = sub[sub['Inception Year'].isin(early_yrs)]
        early_total = len(early_sub)
        early_types = early_sub['产品类型'].value_counts()
        early_desc = "、".join([f"{TYPE_SHORT[t]}{early_types[t]}只" for t in early_types.index[:2]])
        parts.append(f"{min(early_yrs)}-{max(early_yrs)}年为早期布局阶段，共发行{early_total}只产品，以{early_desc}为主")

    # Middle phase
    mid_yrs = [y for y in all_years if y > max(early_yrs) and y <= max(all_years) - 2] if len(all_years) > 5 else []
    if mid_yrs:
        mid_sub = sub[sub['Inception Year'].isin(mid_yrs)]
        mid_total = len(mid_sub)
        mid_types = mid_sub['产品类型'].value_counts()
        mid_desc = "、".join([f"{TYPE_SHORT[t]}{mid_types[t]}只" for t in mid_types.index[:2]])
        parts.append(f"{min(mid_yrs)}-{max(mid_yrs)}年为集中扩张阶段，共发行{mid_total}只，以{mid_desc}为主")

    # Late phase
    late_yrs = [y for y in all_years if y > max(early_yrs) and (not mid_yrs or y > max(mid_yrs))]
    if late_yrs and len(late_yrs) >= 1:
        late_sub = sub[sub['Inception Year'].isin(late_yrs)]
        late_total = len(late_sub)
        if late_total > 0:
            late_types = late_sub['产品类型'].value_counts()
            late_desc = "、".join([f"{TYPE_SHORT[t]}{late_types[t]}只" for t in late_types.index[:2]])
            yr_label = f"{min(late_yrs)}-{max(late_yrs)}年" if len(late_yrs) > 1 else f"{late_yrs[0]}年"
            parts.append(f"{yr_label}发行{late_total}只，以{late_desc}为主")

    # Add key highlights per firm
    highlights = {
        'JPMorgan': "其策略不在于产品数量多，而在于JEPI和JEPQ两只核心产品各自达到400亿美元以上的规模。",
        'First Trust': "2020-2021年为其Buffer产品集中扩张期，此后持续填充不同月份和缓冲深度的产品缺口。",
        'Innovator ETFs': "作为Buffer品类开创者，2016-2019年定义了核心保护区间后进入横向复制阶段，但后期新发产品的边际规模贡献递减。",
        'Neos Funds': "虽为最晚入场者，但以高收入叙事快速切入市场，2024年QQQI发行仅2.5年即达129亿美元。",
        'Global X Funds': "2013年即发行QYLD和XYLD，为covered call ETF先驱，但此后十余年产品线扩展缓慢，未复制早期产品的规模水平。",
    }

    summary = "从发行节奏看，" + "；".join(parts) + "。"
    if firm in highlights:
        summary += highlights[firm]

    return summary


# ============================================================
# CHAPTER 3
# ============================================================
def generate_chapter3(doc, ref, opt, perf_lu):
    add_h(doc, ref, "三、主要产品类型拆解", 1)

    # Compute fee premium lookup
    opt_fee = opt.dropna(subset=['Fee']).copy()
    opt_fee['Bench'] = opt_fee.apply(classify_benchmark, axis=1)

    # ---- 3.1 收益增强型 ----
    add_h(doc, ref, "3.1 收益增强型产品", 2)

    add_p(doc, ref,
        "收益增强型产品是使用期权ETF中规模最大的类别（AUM约1,945亿美元，占衍生策略总规模的30.76%），"
        "以Covered Call、Buy-Write、Put-Write和Premium Income策略为主。"
        "产品通常持有股票或指数相关敞口，同时卖出认购期权或通过ELN、互换等工具嵌入卖权结构，"
        "以权利金和底层资产分红形成现金流。收益来源可拆为四项：标的资产涨跌、成分股分红、"
        "卖出期权权利金、期权平仓或被执行损益。")

    add_p(doc, ref,
        "这类产品更适合震荡市、温和上涨和隐含波动率较高的市场环境。"
        "其核心代价是让渡上行空间，在快速单边上涨行情中可能明显落后于纯指数ETF；"
        "在急跌行情中，权利金只能提供有限缓冲，无法消除权益下跌风险。"
        "正式分析不能只看分派率，还要看底层资产、期权期限、行权价、覆盖比例、再平衡频率和分派来源。")

    # Helper to safely get a product row by ticker
    def get_p(ticker):
        if ticker in perf_lu.index:
            return perf_lu.loc[ticker]
        return None

    # 3.1.1 JEPI
    add_h(doc, ref, "3.1.1 JEPI：主动权益+ELN的混合路径", 3)
    r = get_p('JEPI')
    if r is not None:
        bench = get_benchmark_for_ticker('JEPI')
        bench_fee_str = f"{bench['ticker']} {bench['fee']:.2f}%" if bench else "N/A"
        add_p(doc, ref,
            f"JEPI于2020年5月20日发行，AUM约445.48亿美元，年化费用0.35%。"
            f"产品通过两层结构实现收益目标：第一层是主动管理的美国大盘低波动股票组合；"
            f"第二层是最多20%净资产投资于ELN，ELN嵌入S&P 500的卖出认购期权敞口。"
            f"绩效方面，近1年回报{fmt_r(r['Ret_1Yr'])}，近3年年化{fmt_r(r['Ret_3Yr'])}，"
            f"1年波动率{fmt_r(r['StdDev_1Yr'])}，3年Beta {r['Beta_3Yr']:.2f}，"
            f"近1年Sharpe {r['Sharpe_1Yr']:.2f}。"
            f"相对其底层基准S&P 500纯指数ETF（{bench_fee_str}），JEPI的费率溢价为"
            f"{fmt_bp(0.35 - bench['fee'])}（0.35% vs {bench['fee']:.2f}%），"
            f"反映了主动管理和ELN结构的附加价值。")

    # 3.1.2 JEPQ
    add_h(doc, ref, "3.1.2 JEPQ：Nasdaq成长+期权的收益增强", 3)
    r = get_p('JEPQ')
    if r is not None:
        bench = get_benchmark_for_ticker('JEPQ')
        bench_fee_str = f"{bench['ticker']} {bench['fee']:.2f}%" if bench else "N/A"
        jepi = get_p('JEPI')
        jepi_ret = fmt_r(jepi['Ret_1Yr']) if jepi is not None else "N/A"
        add_p(doc, ref,
            f"JEPQ于2022年5月3日发行，AUM约401.53亿美元，年化费用0.35%。"
            f"与JEPI不同，JEPQ的权益底层围绕Nasdaq 100构建，成长和科技属性更强。"
            f"近1年回报{fmt_r(r['Ret_1Yr'])}（vs JEPI {jepi_ret}），"
            f"1年波动率{fmt_r(r['StdDev_1Yr'])}，3年Beta {r['Beta_3Yr']:.2f}，"
            f"Sharpe Ratio {r['Sharpe_1Yr']:.2f}，"
            f"在承担更高波动的前提下提供了更高的风险调整后收益。"
            f"相对Nasdaq 100纯指数ETF（{bench_fee_str}），费率溢价为{fmt_bp(0.35 - bench['fee'])}。")

    # 3.1.3 Comparison
    add_h(doc, ref, "3.1.3 规则化Buy-Write与高收入Option Income路径对比", 3)

    add_p(doc, ref,
        "在收益增强型产品中，存在两种典型的实现路径："
        "一是以QYLD、XYLD为代表的「规则化buy-write」——持有成熟宽基指数敞口，"
        "系统性卖出指数认购期权，追踪公开策略指数；"
        "二是以SPYI、QQQI为代表的「高收入option income」——通过主动或半主动的期权管理，"
        "在追求更高分派率的同时保持一定的资本增值参与度。"
        "以下对比表加入对应纯指数ETF费率作为参照，以直观展示费率溢价水平。")

    add_ttl(doc, ref, "表：收益增强型代表产品多维度对比（含基准ETF费率）")

    compare = ['JEPI','JEPQ','QYLD','XYLD','SPYI','QQQI']
    comp_rows = []
    for t in compare:
        r = get_p(t)
        if r is None: continue
        bench = get_benchmark_for_ticker(t)
        bench_str = f"{bench['ticker']}（{bench['fee']:.2f}%）" if bench else "N/A"
        fee_prem = fmt_bp(r['Fee'] - bench['fee']) if bench and pd.notna(r['Fee']) else "N/A"
        path = ("主动+ELN" if t in ['JEPI','JEPQ']
                else "规则化buy-write" if t in ['QYLD','XYLD']
                else "active option income")
        comp_rows.append([
            t, str(r.get('Firm Name','')), path,
            fmt_f(r['Fee']), bench_str, fee_prem,
            f"{r['Fund Size USD']/1e8:.2f}" if r['Fund Size USD'] > 0 else "N/A",
            fmt_r(r['Ret_1Yr']), fmt_r(r['Ret_3Yr']),
            fmt_r(r['StdDev_1Yr']),
            f"{r['Sharpe_1Yr']:.2f}" if pd.notna(r['Sharpe_1Yr']) else "N/A",
            f"{r['Beta_1Yr']:.2f}" if pd.notna(r['Beta_1Yr']) else "N/A",
        ])
    add_tbl(doc, ref,
        ["Ticker", "管理人", "策略路径", "费率", "基准ETF（费率）", "费率溢价",
         "AUM（亿美元）", "近1年回报", "近3年年化", "1年波动率", "Sharpe 1Yr", "Beta 1Yr"],
        comp_rows, font_size=7.2)
    add_src(doc, ref,
        "J.P. Morgan、Nasdaq、NEOS、Global X、Morningstar、易方达产品研究。"
        "基准ETF费率来自BlackRock/Invesco/SSGA官网及Nasdaq、Morningstar等公开数据（2026年6月访问）。")

    add_p(doc, ref,
        "从对比可见：规则化buy-write产品（QYLD/XYLD）的Beta最低（0.36-0.45），波动率也最低，"
        "但近1年回报远低于主动管理型产品，反映出在科技股牛市中让渡上行空间的高昂代价。"
        "主动+ELN路径（JEPI/JEPQ）在控制波动的前提下实现了更高的风险调整后收益（Sharpe 0.59-2.17）。"
        "纯active option income路径（SPYI/QQQI）近1年回报最高（24%-31%），但波动率和Beta也更高。"
        "从商业化角度看，三条路径均相对纯指数ETF拥有显著费率溢价——"
        "JEPI/JEPQ相对各自基准溢价20-32bp，SPYI/QQQI溢价53-65bp，QYLD/XYLD溢价45-57bp，"
        "这意味着即使策略规则化程度高、复制门槛低，基金公司仍可获得远高于纯指数ETF的管理费收入。")

    # ---- 3.2 风险缓冲型 ----
    add_h(doc, ref, "3.2 风险缓冲型产品", 2)

    add_p(doc, ref,
        "风险缓冲型产品（Buffer/Target Outcome ETF）通过FLEX Options构建期权组合，"
        "预先设定一段持有期内的风险收益边界：下跌端提供一定比例缓冲（如9%、15%或30%），"
        "上涨端通常设置收益上限（cap）。该类产品AUM约902亿美元（占衍生策略总规模的14.27%），"
        "集中度在四大类别中最低（Top5仅占18.99%），产品形态和策略设计仍有较大差异化空间。")

    # 3.2.1 BUFR
    add_h(doc, ref, "3.2.1 BUFR：Laddered Buffer的常青设计", 3)
    r = get_p('BUFR')
    if r is not None:
        bench = get_benchmark_for_ticker('BUFR')
        bench_str = f"{bench['ticker']} {bench['fee']:.2f}%" if bench else "N/A"
        add_p(doc, ref,
            f"BUFR于2023年5月25日发行，AUM约97.32亿美元，管理费0.10%。"
            f"其核心创新在于Laddered Buffer设计：基金持有四只不同到期月份的底层Target Outcome Buffer ETF，"
            f"每三个月有一只底层ETF进入新的为期一年的目标结果期间，相应刷新buffer和cap。"
            f"这一设计降低了单一建仓时点的影响，使产品更接近长期配置工具。"
            f"绩效方面，近1年回报{fmt_r(r['Ret_1Yr'])}，1年波动率仅{fmt_r(r['StdDev_1Yr'])}，"
            f"Sharpe Ratio {r['Sharpe_1Yr']:.2f}，3年Beta {r['Beta_3Yr']:.2f}，"
            f"体现了Buffer结构在控制下行风险方面的有效性。"
            f"需注意BUFR的0.10%仅为管理费，其底层ETF另收取费用，投资者实际承担的总费率高于此数。")

    # 3.2.2 Innovator Buffer
    add_h(doc, ref, "3.2.2 Innovator Buffer系列：品类定义的范本", 3)

    add_p(doc, ref,
        "Innovator的Buffer系列是理解目标结果产品货架化的最佳范本。"
        "其核心设计逻辑为：先定义保护区间（buffer）、收益上限（cap）和结果期间（outcome period），"
        "再按底层资产、月份和保护强度进行横向复制。")

    add_ttl(doc, ref, "表：Innovator代表Buffer产品绩效对比（含基准ETF费率）")
    buf_prods = ['BALT','BUFF','BUFB']
    buf_rows = []
    for t in buf_prods:
        r = get_p(t)
        if r is None: continue
        bench = get_benchmark_for_ticker(t)
        bench_str = f"{bench['ticker']}（{bench['fee']:.2f}%）" if bench else "N/A"
        fee_prem = fmt_bp(r['Fee'] - bench['fee']) if bench and pd.notna(r['Fee']) else "N/A"
        buf_rows.append([
            t, f"{r['Fund Size USD']/1e8:.2f}", fmt_f(r['Fee']),
            bench_str, fee_prem,
            fmt_r(r['Ret_1Yr']), fmt_r(r['Ret_3Yr']),
            fmt_r(r['StdDev_1Yr']),
            f"{r['Sharpe_1Yr']:.2f}" if pd.notna(r['Sharpe_1Yr']) else "N/A",
            f"{r['Beta_1Yr']:.2f}" if pd.notna(r['Beta_1Yr']) else "N/A",
        ])
    add_tbl(doc, ref,
        ["Ticker", "AUM（亿美元）", "费率", "基准ETF（费率）", "费率溢价",
         "近1年回报", "近3年年化", "1年波动率", "Sharpe 1Yr", "Beta 1Yr"],
        buf_rows, font_size=7.8)
    add_src(doc, ref,
        "Innovator、Morningstar、易方达产品研究。基准ETF费率来源同上。")

    add_p(doc, ref,
        "从产品设计角度看，BALT（3个月20% Buffer）与BUFR（Laddered Buffer）代表了Buffer产品的两种演化方向："
        "前者以更短的outcome period争取更频繁的buffer刷新，适合对建仓时点敏感度较高的投资者；"
        "后者以Laddered结构实现「买入并持有」的配置体验，适合将Buffer作为长期组合组件的投资者。"
        "对国内产品开发而言，Buffer产品的核心挑战在于：FLEX Options在国内尚无对应工具，"
        "标准化场内期权的期限和行权价难以精确复制buffer和cap；投资者教育需前置——"
        "「buffer不是保本」「中途买入的剩余保护区间会变化」「超出缓冲区间的亏损仍由投资者承担」"
        "等概念需要清晰传达。")

    # ---- 3.3 杠杆与反向型 ----
    add_h(doc, ref, "3.3 杠杆与反向型产品", 2)

    add_p(doc, ref,
        "杠杆与反向型ETF通过期货、互换等衍生工具提供标的每日多倍收益或反向收益。"
        "该类产品共1,843只，AUM约3,022亿美元（占衍生策略总规模的47.79%），"
        "但其中仅21只使用期权工具，绝大多数通过期货、互换和每日再平衡实现。"
        "该类产品交易属性强，核心风险来自每日再平衡的波动拖累（volatility decay）、"
        "路径依赖和杠杆放大效应。对国内公募产品开发而言，借鉴价值主要在于理解海外创新边界，"
        "而非直接复制。")

    # ---- 3.4 另类衍生策略型 ----
    add_h(doc, ref, "3.4 另类衍生策略型产品", 2)

    add_p(doc, ref,
        "另类衍生策略型产品覆盖尾部风险对冲、波动率策略、管理期货、多资产叠加和资本效率等多元策略，"
        "共468只产品，AUM约454亿美元，其中115只使用期权工具。该类产品内部差异最大，"
        "应按组合功能而非简单按收益率排序进行拆解。")

    # 3.4.1
    add_h(doc, ref, "3.4.1 波动率与尾部风险策略", 3)

    add_p(doc, ref,
        "波动率和尾部风险产品解决极端市场环境下的风险管理问题。代表产品SVOL、VIXY和TAIL"
        "分别代表波动率风险溢价捕获、VIX期货敞口和尾部保护三种思路。")

    add_ttl(doc, ref, "表：波动率与尾部风险代表产品对比")
    vol_prods = ['SVOL','VIXY','TAIL']
    vol_rows = []
    strat_map = {'SVOL': '做空VIX期货+期权，赚取波动率风险溢价',
                 'VIXY': '做多VIX短期期货，尾部对冲工具',
                 'TAIL': '通过保护性期权管理极端下跌风险'}
    for t in vol_prods:
        r = get_p(t)
        if r is None: continue
        vol_rows.append([
            t, str(r.get('Firm Name','')),
            f"{r['Fund Size USD']/1e8:.2f}" if r['Fund Size USD'] > 0 else "N/A",
            fmt_f(r['Fee']), fmt_r(r['Ret_1Yr']), fmt_r(r['Ret_3Yr']),
            fmt_r(r['StdDev_1Yr']),
            f"{r['Beta_1Yr']:.2f}" if pd.notna(r['Beta_1Yr']) else "N/A",
            strat_map.get(t, ''),
        ])
    add_tbl(doc, ref,
        ["Ticker", "管理人", "AUM（亿美元）", "费率", "近1年回报", "近3年年化",
         "1年波动率", "Beta 1Yr", "策略特征"],
        vol_rows, font_size=7.8)
    add_src(doc, ref, "Simplify、ProShares、Cambria、Morningstar、易方达产品研究")

    add_p(doc, ref,
        "SVOL通过做空VIX期货并叠加期权管理尾部风险，近1年回报12.36%反映了波动率风险溢价"
        "在正常市场环境中的正收益特征，但该策略在波动率急剧飙升时可能面临较大回撤。"
        "TAIL以约0.59%的费率维持保护性期权组合，近1年回报-9.48%，反映牛市中的保护成本——"
        "这类产品的评价标准不应是独立收益率，而应是在组合层面降低最大回撤和尾部相关性的能力。"
        "VIXY直接跟踪VIX短期期货指数，受期货展期损耗（contango drag）影响，"
        "长期持有收益通常为负，更适合作为短期对冲工具而非长期配置。")

    # 3.4.2
    add_h(doc, ref, "3.4.2 管理期货与资本效率策略", 3)

    add_p(doc, ref,
        "管理期货和资本效率产品补充传统股债以外的收益来源，核心价值在于与股债资产的低相关性。"
        "代表产品包括DBMF、CTA（管理期货趋势跟随）和NTSX（资本效率叠加策略）。")

    add_ttl(doc, ref, "表：管理期货与资本效率代表产品对比")
    mf_prods = ['DBMF','CTA','NTSX']
    mf_rows = []
    mf_strat = {'DBMF': '多资产期货趋势跟随，覆盖股票、债券、商品、外汇',
                'CTA': '量化趋势信号驱动的多资产管理期货',
                'NTSX': '90%大盘股+60%国债期货，资本效率叠加'}
    for t in mf_prods:
        r = get_p(t)
        if r is None: continue
        mf_rows.append([
            t, str(r.get('Firm Name','')),
            f"{r['Fund Size USD']/1e8:.2f}" if r['Fund Size USD'] > 0 else "N/A",
            fmt_f(r['Fee']), fmt_r(r['Ret_1Yr']), fmt_r(r['Ret_3Yr']),
            fmt_r(r['StdDev_1Yr']),
            f"{r['Beta_1Yr']:.2f}" if pd.notna(r['Beta_1Yr']) else "N/A",
            mf_strat.get(t, ''),
        ])
    add_tbl(doc, ref,
        ["Ticker", "管理人", "AUM（亿美元）", "费率", "近1年回报", "近3年年化",
         "1年波动率", "Beta 1Yr", "策略特征"],
        mf_rows, font_size=7.8)
    add_src(doc, ref, "iMGP/DBi、Simplify、WisdomTree、Morningstar、易方达产品研究")

    add_p(doc, ref,
        "DBMF是规模最大的管理期货ETF（AUM 40.35亿美元），近1年回报30.05%，Beta仅0.30，"
        "体现了管理期货策略与股债资产在波动期的低相关优势。"
        "CTA采用类似趋势跟随逻辑，3年Beta为-0.21，与传统资产呈微弱负相关。"
        "NTSX以「90/60」资本效率结构著称——约90%资产投资于美国大盘股票，"
        "同时通过国债期货获得约60%的债券敞口，在同一资本基础上叠加股债两类资产暴露。"
        "国内借鉴时应优先研究其资产配置功能：管理期货策略公募化需解决策略透明度和流动性管理问题；"
        "资本效率策略受衍生品净敞口限制，制度突破前难以直接复制。")

    # 3.4.3
    add_h(doc, ref, "3.4.3 现金替代与混合型界定收益策略", 3)

    add_p(doc, ref,
        "此类产品通过期权结构或固收底仓改善收益路径，代表产品包括BOXX"
        "（通过box spreads实现类现金收益+税务优势）和各类option collar策略产品。"
        "BOXX利用期权盒式价差构建具有税务效率的现金替代工具，在应税账户中具有独特的配置价值。"
        "目前Morningstar对BOXX的覆盖数据有限（费率和AUM均待补充），"
        "但其代表的「通过期权结构优化现金管理」思路，对国内低利率环境下的现金管理类产品创新具有启发意义。"
        "需要特别提醒的是，现金替代类期权策略的税务处理、估值透明度和投资者理解难度是制约其推广的三大障碍。")

    # ---- 3.5 Fee premium summary ----
    add_h(doc, ref, "3.5 费率溢价与商业化价值小结", 2)

    add_p(doc, ref,
        "综合以上四类产品的费率对比可以看出，期权策略ETF相对纯指数ETF的费率溢价"
        "是其核心商业价值之一。从底层资产维度看：S&P 500类期权ETF的AUM加权费率为0.41%"
        "（相对IVV 0.03%溢价38bp），Nasdaq 100类为0.47%（相对QQQM 0.15%溢价32bp），"
        "Russell 2000类为0.62%（相对IWM 0.19%溢价43bp）。")

    add_p(doc, ref,
        "将期权策略ETF置于更广义的基金费率谱系中，其商业定位更为清晰。"
        "根据ICI和Morningstar发布的2025年度基金费用研究："
        "美国指数ETF的资产加权平均费率为0.05%-0.14%，"
        "美国主动管理ETF的简单平均费率约0.43%-0.44%，"
        "美国主动权益基金整体的资产加权平均费率为0.58%。"
        "期权策略ETF的AUM加权费率恰好处于「指数之上、主动之下」的中间地带——"
        "在维持ETF结构透明性和流动性的同时，实现了接近主动管理基金的费率水平。")

    add_p(doc, ref,
        "对国内基金公司而言，这意味着期权策略ETF是一件费率与规模双赢的战略选择："
        "相比纯指数ETF，有更高的单位AUM管理费收入；相比主动管理基金，有更低的运营成本"
        "和更好的可扩展性。策略复杂度与费率溢价正相关的规律——Buffer和高收入option income"
        "策略可获得最高的费率溢价（50-65bp）——也为基金公司提供了清晰的战略方向："
        "沿着「规则化→主动化→结构化」的路径逐步提升产品附加值和费率水平。")


def get_benchmark_for_ticker(ticker):
    """Get benchmark info for a ticker"""
    mapping = {
        'JEPI': 'S&P 500', 'XYLD': 'S&P 500', 'SPYI': 'S&P 500',
        'HELO': 'S&P 500', 'BALT': 'S&P 500', 'BUFR': 'S&P 500',
        'JEPQ': 'Nasdaq 100', 'QYLD': 'Nasdaq 100', 'QQQI': 'Nasdaq 100',
        'RYLD': 'Russell 2000', 'DIVO': 'Dow Jones',
    }
    bench_key = mapping.get(ticker)
    return BENCHMARKS.get(bench_key) if bench_key else None


# ============================================================
# WORKBOOK UPDATE
# ============================================================
def update_workbook(opt, mgr, perf_lu):
    wb = load_workbook(WB_PATH)

    thin = Side(style="thin", color=BORDER_COLOR)
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def style_sheet(ws):
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.font = Font(name=CN_FONT, size=10)
                cell.border = border
        for cell in ws[1]:
            cell.fill = PatternFill("solid", fgColor=HEADER_FILL)
            cell.font = Font(name=CN_FONT, size=10, bold=True)

    # Existing sheets to replace
    for sname in ["管理人排行", "产品发行时间线", "资金流向分析", "AUM积累效率",
                  "管理人产品类型年份交叉", "产品绩效补全", "费率溢价分析"]:
        if sname in wb.sheetnames: del wb[sname]

    # Sheet: 管理人排行
    ws_mgr = wb.create_sheet("管理人排行")
    ws_mgr.append(["排名", "管理人", "产品数量", "总AUM（亿美元）", "AUM加权费率",
                   "收益增强型", "风险缓冲型", "杠杆与反向型", "另类衍生策略型"])
    for idx, (_, r) in enumerate(mgr.head(30).iterrows()):
        ws_mgr.append([idx+1, r['Firm Name'], int(r['产品数量']), round(r['总AUM亿'], 2),
                       round(r.get('AUM加权费率', 0), 4) if pd.notna(r.get('AUM加权费率')) else None,
                       int(r.get('收益增强型_数量', 0)), int(r.get('风险缓冲型_数量', 0)),
                       int(r.get('杠杆与反向型_数量', 0)), int(r.get('另类衍生策略型_数量', 0))])
    style_sheet(ws_mgr)

    # Sheet: 产品发行时间线 (year × type)
    ws_tl = wb.create_sheet("产品发行时间线")
    ws_tl.append(["管理人", "年份", "产品类型", "新增产品数"])
    for firm in mgr.head(5)['Firm Name']:
        sub = opt[opt['Firm Name'] == firm]
        for yr in sorted(sub['Inception Year'].unique()):
            yr_data = sub[sub['Inception Year'] == yr]
            for typ in TYPE_ORDER:
                cnt = int((yr_data['产品类型'] == typ).sum())
                if cnt > 0:
                    ws_tl.append([firm, int(yr), typ, cnt])
    style_sheet(ws_tl)

    # Sheet: 资金流向分析
    ws_flow = wb.create_sheet("资金流向分析")
    ws_flow.append(["管理人", "时间窗口", "净流量（亿美元）", "流入产品数", "流出产品数",
                    "缺失产品数", "流入合计（亿美元）", "流出合计（亿美元）", "最大贡献产品"])
    flow_labels = [('Flow_1Mo','1个月'), ('Flow_3Mo','3个月'), ('Flow_6Mo','6个月'),
                   ('Flow_YTD','YTD'), ('Flow_1Yr','1年'), ('Flow_3Yr','3年')]
    for firm in mgr.head(30)['Firm Name']:
        sub = opt[opt['Firm Name'] == firm]
        for fl_col, fl_label in flow_labels:
            has_data = sub[sub[fl_col].notna()]
            n_data = len(has_data)
            n_missing = len(sub) - n_data
            total = has_data[fl_col].sum() / 1e8
            n_in = int((has_data[fl_col] > 0).sum())
            n_out = int((has_data[fl_col] < 0).sum())
            inflow_sum = has_data[has_data[fl_col] > 0][fl_col].sum() / 1e8
            outflow_sum = has_data[has_data[fl_col] < 0][fl_col].sum() / 1e8
            top_str = ""
            if n_data > 0:
                top_idx = has_data[fl_col].abs().nlargest(1).index[0]
                top_str = f"{has_data.loc[top_idx, 'Ticker']}（{has_data.loc[top_idx, fl_col]/1e8:+.2f}亿）"
            ws_flow.append([firm, fl_label, round(total, 2), n_in, n_out, n_missing,
                           round(inflow_sum, 2), round(outflow_sum, 2), top_str])
    style_sheet(ws_flow)

    # Sheet: 产品绩效补全
    ws_perf = wb.create_sheet("产品绩效补全")
    ws_perf.append(["Ticker", "产品名称", "管理人", "成立年份", "AUM（亿美元）", "费率",
                    "近1年回报", "近3年年化", "1年波动率", "3年波动率",
                    "1年Sharpe", "3年Sharpe", "1年Beta", "3年Beta",
                    "1月净流（亿）", "1年净流（亿）", "3年净流（亿）"])
    for t, r in perf_lu.iterrows():
        ws_perf.append([
            t, r.get('Name',''), str(r.get('Firm Name','')),
            int(r['Inception Year']) if pd.notna(r.get('Inception Year')) else None,
            round(r['Fund Size USD']/1e8, 2) if r['Fund Size USD'] > 0 else None,
            round(r['Fee'], 4) if pd.notna(r.get('Fee')) else None,
            round(r['Ret_1Yr'], 2) if pd.notna(r.get('Ret_1Yr')) else None,
            round(r['Ret_3Yr'], 2) if pd.notna(r.get('Ret_3Yr')) else None,
            round(r['StdDev_1Yr'], 2) if pd.notna(r.get('StdDev_1Yr')) else None,
            round(r['StdDev_3Yr'], 2) if pd.notna(r.get('StdDev_3Yr')) else None,
            round(r['Sharpe_1Yr'], 2) if pd.notna(r.get('Sharpe_1Yr')) else None,
            round(r['Sharpe_3Yr'], 2) if pd.notna(r.get('Sharpe_3Yr')) else None,
            round(r['Beta_1Yr'], 2) if pd.notna(r.get('Beta_1Yr')) else None,
            round(r['Beta_3Yr'], 2) if pd.notna(r.get('Beta_3Yr')) else None,
            round(r.get('Flow_1Mo', 0)/1e8, 2) if pd.notna(r.get('Flow_1Mo')) and r.get('Flow_1Mo', 0) != 0 else None,
            round(r.get('Flow_1Yr', 0)/1e8, 2) if pd.notna(r.get('Flow_1Yr')) and r.get('Flow_1Yr', 0) != 0 else None,
            round(r.get('Flow_3Yr', 0)/1e8, 2) if pd.notna(r.get('Flow_3Yr')) and r.get('Flow_3Yr', 0) != 0 else None,
        ])
    style_sheet(ws_perf)

    # Update source list
    src_ws = wb["来源清单"]
    existing = set()
    for row in range(2, src_ws.max_row + 1):
        v = src_ws.cell(row, 3).value
        if v: existing.add(str(v).strip())
    new_srcs = [
        ("S33","网页","iShares Core S&P 500 ETF (IVV)","https://www.blackrock.com/us/individual/products/239726/","2026-06-18","Ch3费率对比","0.03%"),
        ("S34","网页","Invesco Nasdaq 100 ETF (QQQM)","https://www.invesco.com/us/financial-products/etfs/product-detail?ticker=QQQM","2026-06-18","Ch3费率对比","0.15%"),
        ("S35","网页","iShares Russell 2000 ETF (IWM)","https://www.ishares.com/us/products/239710/","2026-06-18","Ch3费率对比","0.19%"),
        ("S36","网页","SPDR Dow Jones ETF Trust (DIA)","https://www.ssga.com/us/en/intermediary/etfs/spdr-dow-jones-industrial-average-etf-trust-DIA","2026-06-18","Ch3费率对比","0.16%"),
        ("S37","报告","ICI 2025年度基金费用报告","https://www.ici.org/news-release/mutual-fund-and-etf-fees-remained-near-historic-lows-in-2025","2025","Ch3行业费率锚点","指数/主动基金费率"),
        ("S38","报告","Morningstar 2026年度基金费用研究","https://www.morningstar.com/funds/how-active-etfs-are-reshaping-fund-fees","2026","Ch3行业费率锚点","主动ETF费率"),
    ]
    for sid, stype, title, link, date, usage, status in new_srcs:
        if title not in existing:
            src_ws.append([sid, stype, title, link, date, usage, status])
    style_sheet(src_ws)

    wb.save(OUT_WB)
    print(f"Workbook saved to: {OUT_WB}")


# ============================================================
# MAIN
# ============================================================
def main():
    print("=" * 60)
    print("Ch2&3 重写脚本 V2")
    print("=" * 60)

    print("[1/4] 加载数据...")
    raw = load_data()
    opt = get_opt(raw)
    # Normalize Firm Name to avoid splitting managers across name variants
    firm_name_map = {
        'JPMorgan': ['JPMorgan', 'JPMorgan Chase'],
        'First Trust': ['First Trust'],
        'Innovator ETFs': ['Innovator', 'Innovator ETFs'],
        'Neos Funds': ['Neos', 'Neos Funds'],
        'Global X Funds': ['Global X', 'Global X Funds', 'Global X Investments Canada Inc.'],
    }
    def normalize_firm(name):
        name_s = str(name)
        for canonical, variants in firm_name_map.items():
            for v in variants:
                if v in name_s:
                    return canonical
        return name_s
    opt['Firm Name'] = opt['Firm Name'].apply(normalize_firm)
    raw['Firm Name'] = raw['Firm Name'].apply(normalize_firm)
    # Dedup by ticker to avoid double-counting share classes
    opt = opt.sort_values('Fund Size USD', ascending=False).drop_duplicates(subset=['Ticker'], keep='first')
    perf_lu = get_perf(raw)
    print(f"  使用期权产品: {len(opt)} 只, AUM ${opt['Fund Size USD'].sum()/1e8:.2f}亿")

    print("[2/4] 构建分析...")
    mgr = build_manager_ranking(opt)
    for _, r in mgr.head(5).iterrows():
        print(f"  {r['Firm Name']}: {int(r['产品数量'])}只, AUM ${r['总AUM亿']:.2f}亿")

    print("[3/4] 更新Word...")
    doc = Document(str(DOC_PATH))
    ref = remove_between_headings(doc, "二、", "四、")
    generate_chapter2(doc, ref, opt, mgr, perf_lu)
    generate_chapter3(doc, ref, opt, perf_lu)
    doc.save(str(OUT_DOC))
    print(f"  Word: {OUT_DOC}")

    print("[4/4] 更新底稿...")
    update_workbook(opt, mgr, perf_lu)

    print("\n完成!")

if __name__ == "__main__":
    main()
